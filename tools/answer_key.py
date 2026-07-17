#!/usr/bin/env python3
"""Build the Answer_Key folder: data inconsistencies and expected deliverables."""

from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

from brand import (
    add_body,
    add_bullet,
    add_cover,
    add_heading,
    add_number,
    add_table,
    set_doc_defaults,
)

ROOT = Path(__file__).resolve().parents[1]
KEY = ROOT / "Answer_Key"

NAVY = "1B2A4A"
TEAL = "007A87"
WHITE = "FFFFFF"

# ---------------------------------------------------------------------------
# Dataset inconsistencies (built into generate_data.py)
# ---------------------------------------------------------------------------

DATASET_ISSUES = [
    ("DC-01", "Duplicate beneficiary IDs", "beneficiary_registry",
     "Two IDs appear twice with conflicting province values and different last_updated dates.",
     "De-duplicate before any count; keep the latest last_updated and log the conflict."),
    ("DC-02", "Mixed date formats", "All tables with dates",
     "Same field uses dd/mm/yyyy, yyyy-mm-dd and dd-mm-yyyy interchangeably.",
     "Standardise to one format before joins; document the rule in the assumption log."),
    ("DC-03", "Inconsistent placement coding", "employment_outcomes",
     "placed_90d uses Y, Yes, 1, N, No, 0 and blank.",
     "Map to a single binary flag; blank = missing follow-up, not 'not placed'. "
     "Primary VPBank rate uses completion_status = Completed only."),
    ("DC-04", "Inconsistent attendance coding", "programme_attendance",
     "attended uses Y, N, 1 and 0. BM uses fewer mentor-circle sessions than SF/PD classroom sessions.",
     "Normalise to 1/0; analyse BM separately or with a session-count caveat."),
    ("DC-05", "Impossible age", "beneficiary_registry",
     "Record GBF-2025-10417 has a date of birth implying age ~42 (one intentional error).",
     "Flag as data entry error; exclude from age analysis; do not treat the whole file as out-of-eligibility."),
    ("DC-06", "Wage typo", "employment_outcomes",
     "Record GBF-2025-10417 shows wage_vnd_monthly = 99,999,999 (notes field flags error).",
     "Investigate; exclude or cap for wage averages; never report uncorrected."),
    ("DC-07", "Missing ethnicity", "beneficiary_registry",
     "Roughly 18% of ethnicity fields are blank.",
     "Report missingness; do not impute without stating why."),
    ("DC-08", "Missing disability flag", "beneficiary_registry",
     "Roughly 5% of disability_status fields are blank.",
     "Treat as unknown; compare Y-only rate with and without unknowns."),
    ("DC-09", "Province label mismatch", "beneficiary_registry",
     "Same geography appears as 'Ho Chi Minh City' and 'HCMC'.",
     "Create a lookup table; never double-count provinces."),
    ("DC-10", "Dong Nai cost double-entry", "hub_costs_2025",
     "September 2025 for hub DN appears twice with identical totals.",
     "Remove duplicate row before summing hub costs."),
    ("DC-11", "Hours on inactive mentor", "volunteer_hours",
     "Mentor M-099 has status=Inactive but still has hours_logged.",
     "Filter status=Active for mentor-hour averages, or justify including Inactive hours."),
    ("DC-12", "Text-formatted funding amount", "funding_by_source",
     "UK aid / FCDO 2025-Q3 amount_vnd_m is stored as text with a thousands comma (e.g. '1,250').",
     "Convert to number before SUM; note in data quality report."),
    ("DC-13", "Export aligns with programme headcount", "beneficiary_registry",
     "Registry is sized to ~800 (programme guideline: served in 2025). Waitlist (~1,100) is a separate snapshot. "
     "D-01 reports 2024 completers (592) and public 71% placement - a different grain.",
     "Do not add waitlist to served totals. Do not equate D-01 completers with registry rows. "
     "State placement definition whenever reporting rates."),
    ("DC-14", "Capacity headroom for scale", "hub_capacity_jan2026",
     "DN/LA show overloaded caseloads and thin staff/mentor benches relative to throughput caps.",
     "Funding scenarios and growth options must stress-test staff, mentor and finance capacity - not cost alone."),
]

DOCUMENT_ISSUES = [
    ("DC-D01", "Placement rate definition", "D-01 vs D-02 vs ST-02",
     "D-01 public report: 71% within 90 days (broad: includes gig work >=15 hrs/week). "
     "ST-02: formal contract rate closer to 63%. D-02 says Pathway Digital reporting was generous.",
     "Teams must state which definition they use and report both where possible."),
    ("DC-D02", "Pathway Digital completion", "D-02 vs D-01",
     "Donor-facing figure 68%; true completers/ starters = 52% (wrong denominator).",
     "Strong teams catch this and discuss reputational risk."),
    ("DC-D03", "Dong Nai performance", "D-02, ST-01, ST-05, EP-02",
     "Internship drop-out 38% vs org average 22%; transport and employer quality cited.",
     "Dashboard should show DN below 65%; recommendations must address DN if scaling."),
    ("DC-D04", "Disability gap", "ST-03 vs D-01",
     "Reported disability share ~6-7%; population nearer 9%. Placement ~45-50% vs ~70% on the public broad definition.",
     "Mark down if disability breakdown is missing from analysis or recommendations."),
    ("DC-D05", "Board email thread", "D-08",
     "Forwarded anonymously; unverified; politically motivated opinions.",
     "Use for stakeholder sentiment only; triangulate with transcripts."),
    ("DC-D06", "YouthWorks threat", "D-10 vs ST-04 vs ST-05",
     "Development says overstated; field staff see waitlist leakage in Binh Duong border districts.",
     "No single truth; evidence-based judgement required."),
    ("DC-D07", "VPBank conditions", "D-03 vs DN-01",
     "65% at 90 days required; DN-01 open to quality-focused alternative but not open-ended delay.",
     "Funding scenarios must stress-test the 65% condition honestly."),
    ("DC-D08", "Shortening the placement window", "ST-04 vs D-03 vs DN-01",
     "Anh Duc pushes 'earlier management visibility' / leading indicators without naming a "
     "new primary metric. D-03 Clause 3-4 and DN-01 lock 65% at ninety days; earlier indicators "
     "may be internal only and do not replace disbursement measurement.",
     "Red line if a team elevates a shorter window (e.g. 60-day) as the primary Board/donor "
     "metric to look better. Strong teams catch this by reading the grant, not by being told."),
    ("DC-D09", "Long An pilot", "D-02, ST-01, BD-01",
     "~60% drop-out; board member T.N. pushes expansion for political reasons.",
     "Scaling to Long An without fixing the model is a common weak answer."),
    ("DC-D10", "Social enterprise conflict", "D-12, EP-01",
     "GBF Staffing Solutions earns placement fees from same employers who take interns.",
     "Strong teams surface the conflict in trade-off reflection or options; it is not labelled for them."),
]

DELIVERABLES = [
    ("Project Charter", "All", "Wed 29 Jul", "PDF + Excel RACI",
     "Team names, roles, scope, A↔B evidence loop agreed, online confirm by deadline.",
     "Charter confirmed; RACI has one accountable owner; Estate→Demand→Plan→R/FM rule stated.",
     "Missing confirmations; vague scope; no named WF-A/WF-B leads."),
    ("A1 - Problem statement and issue tree", "A", "Wed 5 Aug", "PDF",
     "One precise decision sentence in the team's own words; MECE branches they invent; 4+ hypotheses that imply asks for B.",
     "Hypotheses map to Analysis Requests; branches emerge from evidence; statement sharpens the engagement letter.",
     "Copies engagement-letter wording; generic or pre-labelled tree; no implied asks."),
    ("A1b - Stakeholder map", "A", "Thu 6 Aug", "PDF",
     "Power/interest grid; five starred stakeholders with cited sources.",
     "Stars match interview priority (VPBank, field staff, beneficiaries, board faction, employers).",
     "Generic list; no sources; treats D-08 as verified fact."),
    ("Credibility matrix", "All", "Thu 6 Aug", "PDF",
     "Every shared client document (D-01, D-02, D-03, D-08, D-10, D-11, D-12) rated H/M/L with justification; Findings Memos rated once received.",
     "Ratings differ by source type; justifications cite audience/incentive/definitions; contradictions noted.",
     "All sources rated 'high' with no justification; copied from a facilitator hint."),
    ("Data Estate Brief", "B", "Mon Week 3", "PDF",
     "Catalog of tables, grain, coverage and limits in plain language. No workbook, dictionary or field dump.",
     "A can aim Demand without seeing raw data; limits are honest; catalog not a spoiler of findings.",
     "Raw file shared; schema dump that lets A prescribe joins; empty or late brief."),
    ("Demand Brief", "A", "Mon-Tue Week 3", "PDF",
     "Ranked decision needs from A1 hypotheses after reading DEB; useful answer shapes; not column shopping.",
     "Top demands map to R-01+; each would change the recommendation if answered.",
     "Vague 'send us the data'; field names invented; filed before reading Estate Brief."),
    ("Analysis Plans", "B", "From Week 3", "PDF",
     "B-owned approach for each priority Demand / R-xx: method, definitions, risks, ETA.",
     "Plan answers A's decision question; A acknowledges without co-editing pivots.",
     "No plan; plan is a field list for A to run; interim raw extracts shared."),
    ("Analysis Requests (min. 3)", "A", "Weeks 3-5", "PDF",
     "Each request states decision, hypothesis, useful answer shape, urgency; linked to Demand/Plan; logged on Request Log.",
     "Asks are specific enough for B to answer; emerge from Demand + transcripts/hypotheses.",
     "Fewer than three; vague 'send us the data'; field names invented from nowhere; no Demand/Plan trail."),
    ("B1 - Data quality report", "B", "Tue 11 Aug", "PDF",
     "DC-01 to DC-14 found or consciously checked; cleaning logged; A gets only material issues via FM.",
     "Duplicates (via last_updated), coding, typo wage, DN double-entry, inactive mentor, text funding row addressed.",
     "Silent cleaning; raw file shared with A; claims 100% clean data."),
    ("Findings Memos (min. 3)", "B", "Weeks 3-5", "PDF",
     "Plain-language answers to R-xx; definitions and caveats; curated charts only.",
     "Restates question; translates honestly; states placement definition and capacity caveats.",
     "Pivot dump; no caveats; raw workbook attached."),
    ("A2 - Interview synthesis", "A", "Wed 12 Aug", "PDF",
     "Thematic summary with transcript refs; contradictions flagged; questions for B listed.",
     "Themes: transport, employer quality, metric integrity, scale vs quality.",
     "Quote dump by person; no contradictions; no citations."),
    ("A3 - Empathy and journey maps", "A", "Wed 12 Aug", "PDF",
     "One archetype grounded in BN-xx; pain points with quotes or FM cites.",
     "Internship stage shows transport cost and employer exploitation; links to ST-05 and EP-02.",
     "Generic persona; no quotes; outside-in design with no lived detail."),
    ("B2 - Performance dashboard", "B", "Fri 14 Aug", "Excel",
     "Placement by hub/programme/gender; flag below 65%; insight-titled charts; B working file only.",
     "DN and Long An flagged; formal vs broad rate shown or footnoted.",
     "Shared raw with A; DN hidden; single blended rate with no definition."),
    ("B3 - Analysis summary", "B", "Wed 19 Aug", "PDF",
     "Headline findings linked to Request/FM IDs; confounders named.",
     "Attendance correlates with placement but transport/motivation confound.",
     "Claims attendance causes placement; ignores confounders."),
    ("A4 - Strategic options", "A", "Thu 20 Aug", "PDF",
     "2-3 distinct options; every number cites R-/FM-; winners/losers named.",
     "Options differ on scale timing; change-my-mind trigger stated.",
     "Three variants of 'scale more'; numbers without Request IDs."),
    ("B4 - Evidence-check memo", "B", "Fri 21 Aug", "PDF",
     "Each A option tested against Findings and D-12/D-03 feasibility.",
     "Flags if aggressive growth cannot hold 65% without selection pressure.",
     "Rubber-stamps A without challenge."),
    ("A5 - Trade-off reflection", "A", "Wed 26 Aug", "PDF",
     "Open prompts answered; who gains/loses; cost of preferred path named.",
     "Team names the operational cost of the ethical choice, or the moral cost of the efficient choice.",
     "Empty platitudes; no named cost; checklist of invented ED codes."),
    ("A6 - Validation protocol", "All", "Wed 26 Aug", "PDF",
     "6+ citations including BN/ST/DN/EP and at least two R-/FM-.",
     "Dong Nai recommendation cites ST-05, BN-02, EP-02 and FM findings.",
     "Validation after the fact; fewer than six citations."),
    ("B5 - Funding scenarios", "B", "Thu 27 Aug", "Excel",
     "Hold / +30% / +80% using D-12, D-03 and hub_capacity_jan2026.",
     "Aggressive scenario shows staffing/mentor gap and 65% at risk; ties choice to mission + Board ambition.",
     "Assumes 65% holds automatically; ignores staff/mentor/finance capacity."),
    ("A7 - Revision log", "All", "Thu 27 Aug", "PDF",
     "At least one substantive change after a B challenge or Findings Memo.",
     "Example: dropped Long An expansion after FM on DN placement.",
     "Empty log or cosmetic edits only."),
    ("Executive summary", "All", "Fri 28 Aug", "PDF",
     "One-sentence key message; recommendation; 90-day roadmap; honest limitations.",
     "Board-ready; cites strongest evidence including Request IDs.",
     "Buries recommendation; no limitations."),
    ("Final report", "All", "Fri 11 Sep", "PDF",
     "15-25 pages; A↔B evidence loop described; every claim sourced.",
     "Integrates A narrative and B findings; trade-offs and risks included.",
     "A and B contradict; unsourced numbers; raw data pasted."),
    ("Board presentation", "All", "Thu 10 Sep", "PowerPoint",
     "Answer-first titles; Dong Nai and formal rate not hidden; cost of path named.",
     "Opens with decision needed; ends with specific Board asks.",
     "Data dump; no clear ask; A cites 'the dataset' without FM IDs."),
    ("Handoff checklist", "All", "Wed 2 Sep", "PDF",
     "Copy 4_Shared_Toolkit/Templates/Cross_Workflow_Handoff_Checklist.docx to team Drive. "
     "Section A (A reviews B), Section B (B reviews A), Section C (Engagement Lead). "
     "Online ticks only: [x] + typed name + YYYY-MM-DD + channel. Export Team_All_Handoff_vN.pdf; "
     "Request Log shows min. three completed R/FM pairs. See Guideline Section 8.",
     "All three sections confirmed; Drive links to A_B_Exchange/ and Handoff PDF present; "
     "genuine exchange (not a last-minute dump).",
     "Missing PDF; wet-signature / print culture; unchecked boxes; fewer than three R/FM pairs; "
     "raw file shared with A."),
    ("Peer evaluation", "Individual", "Mon 14 Sep", "PDF",
     "Confidential peer ratings for each teammate; submitted individually via portal.",
     "Honest differentiation; comments specific to contribution.",
     "Identical scores for everyone; blank comments; late or missing."),
    ("Reflection journal", "Individual", "Tue 15 Sep", "PDF",
     "Personal learning reflection on the engagement and cross-workflow practice.",
     "Names a concrete skill and a cost of the team's recommendation.",
     "Generic essay; no link to the case evidence or teamwork."),
]


def _newdoc(title, subtitle):
    doc = Document()
    set_doc_defaults(doc)
    add_cover(
        doc, title, subtitle,
        classification="ANSWER KEY - FACILITATOR ONLY",
        doc_ref="GBF Case Key",
    )
    return doc


def build_index(docx_to_pdf):
    doc = _newdoc(
        "Answer Key - Start Here",
        "Facilitator reference for data inconsistencies and expected deliverables",
    )
    add_body(
        doc,
        "This folder is for facilitators, buddies and assessors only. Do not distribute "
        "to participants. It consolidates the case key: data inconsistencies (Workflow B), "
        "document contradictions, expected deliverables under the A↔B evidence loop "
        "(Estate → Demand → Plan → R/FM), and solution paths.",
    )
    add_heading(doc, "Lenses we are testing (never tell participants)", 1)
    add_bullet(doc, "Efficiency versus ethics: growth, targets and placement metrics against "
               "dignity, quality and who gets left behind.")
    add_bullet(doc, "Empathy versus outside-in: recommendations grounded in beneficiary voices "
               "versus solving for people whose context the team has not taken seriously.")
    add_bullet(doc, "Active querying: B briefs the data estate; A states demand; B owns the "
               "analysis plan; A learns material inconsistencies only through Findings Memos "
               "(correct by design).")
    add_heading(doc, "Files in this folder", 1)
    add_table(
        doc, ["File", "Purpose"],
        [
            ("01_Data_Inconsistencies_Register.pdf", "Every dataset and document inconsistency"),
            ("02_Expected_Deliverables_and_Outcomes.pdf", "What excellent work looks like per deliverable"),
            ("03_Strategic_Solution_Paths.pdf", "Defensible recommendation paths and red lines"),
            ("Data_Inconsistencies_Register.xlsx", "Searchable register for marking and debrief"),
        ],
        widths=[3.2, 3.2],
    )
    add_heading(doc, "How to use this pack", 1)
    add_number(doc, "Before Week 1: read the data inconsistencies register so you know what teams should find.")
    add_number(doc, "During marking: use Expected Deliverables for criterion-level judgement "
               "(including the online Handoff Checklist standard).")
    add_number(doc, "In debrief: reveal that multiple solution paths were valid (see document 03).")
    add_number(doc, "Operational facilitation remains in Facilitators/ (buddy hints, rubric, "
               "kickoff logistics, debrief flow). Do not distribute Answer_Key or Facilitators "
               "to participants.")

    path = KEY / "Word" / "00_Answer_Key_Index.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    docx_to_pdf(path, KEY / "00_Answer_Key_Index.pdf")
    print("  Answer Key: 00_Answer_Key_Index")


def build_inconsistencies_register(docx_to_pdf):
    doc = _newdoc(
        "Data Inconsistencies Register",
        "Complete list of built-in errors and document contradictions",
    )
    add_body(
        doc,
        "Teams earn credit for finding these, documenting their response, and triangulating "
        "with client documents and transcripts. None should be ignored.",
    )

    add_heading(doc, "Part 1 - Dataset inconsistencies (GBF_Datasets.xlsx - Workflow B only)", 1)
    add_body(doc, "A should only learn material issues through Findings Memos. Mark down if A "
             "opens the raw file or if B dumps the cleaning log without translation.")
    add_table(
        doc, ["Ref", "Issue", "Table", "What teams should see", "Expected handling"],
        [(r[0], r[1], r[2], r[3], r[4]) for r in DATASET_ISSUES],
        widths=[0.6, 1.4, 1.2, 2.0, 1.4],
    )

    add_heading(doc, "Part 2 - Document and narrative contradictions", 1)
    add_table(
        doc, ["Ref", "Topic", "Sources", "Contradiction", "What strong teams do"],
        [(r[0], r[1], r[2], r[3], r[4]) for r in DOCUMENT_ISSUES],
        widths=[0.6, 1.2, 1.2, 2.0, 1.4],
    )

    add_heading(doc, "Part 3 - Expected quantitative patterns (after cleaning)", 1)
    add_bullet(doc, "Among Completed records, HCMC hubs can clear ~65%+ broad; Dong Nai stays below (especially Formal).")
    add_bullet(doc, "Including Dropped in the denominator pulls most hubs under 65% - definition choice matters.")
    add_bullet(doc, "Long An and Pathway Digital underperform Skills Forward on placement.")
    add_bullet(doc, "Broad placement rate sits above formal-only rate (aligns with ST-02 ~63% formal narrative; 2025 export may differ slightly).")
    add_bullet(doc, "Disability placement gap of ~20+ points vs overall (aligns with ST-03).")
    add_bullet(doc, "Higher attendance associated with placement, but confounded by transport and motivation.")
    add_bullet(doc, "Aggressive growth scenario struggles to meet 65% without selection bias "
               "(efficiency vs ethics tension) and fails capacity headroom at DN/LA.")
    add_bullet(doc, "Mentor counts differ by source (D-01 ~95 org-wide active mentors; "
               "hub_capacity_jan2026 active_mentors sum ~26; volunteer_hours has ~44 active IDs "
               "plus Inactive M-099) - strong teams state the grain.")

    path = KEY / "Word" / "01_Data_Inconsistencies_Register.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    docx_to_pdf(path, KEY / "01_Data_Inconsistencies_Register.pdf")
    print("  Answer Key: 01_Data_Inconsistencies_Register")


def build_expected_deliverables(docx_to_pdf):
    doc = _newdoc(
        "Expected Deliverables and Outcomes",
        "Clear standards for each submission",
    )
    add_body(
        doc,
        "Use this document when marking or giving feedback. 'Excellent' is the bar for top "
        "marks; 'Minimum' is pass-level; teams below minimum on core deliverables should "
        "receive targeted buddy support.",
    )

    add_heading(doc, "Summary table", 1)
    add_table(
        doc, ["Deliverable", "WS", "Due", "Format", "Minimum acceptable", "Excellent"],
        [(d[0], d[1], d[2], d[3], d[5], d[6]) for d in DELIVERABLES],
        widths=[1.8, 0.5, 0.8, 0.7, 1.8, 1.8],
    )

    add_heading(doc, "Detailed requirements by deliverable", 1)
    for name, ws, due, fmt, requirement, excellent, minimum_fail in DELIVERABLES:
        add_heading(doc, f"{name} ({ws}, due {due})", 2)
        add_body(doc, "Requirement: " + requirement)
        add_body(doc, "Excellent: " + excellent, bold=True)
        add_body(doc, "Mark down if: " + minimum_fail, italic=True)

    add_heading(doc, "Board presentation - non-negotiables", 1)
    add_bullet(doc, "States placement definition used and shows formal vs broad if both discussed.")
    add_bullet(doc, "Names Dong Nai underperformance or explains why recommendation still scales there.")
    add_bullet(doc, "Engages at least one ethical trade-off with a clear position.")
    add_bullet(doc, "Ends with specific decisions requested from the Board.")

    path = KEY / "Word" / "02_Expected_Deliverables_and_Outcomes.docx"
    doc.save(path)
    docx_to_pdf(path, KEY / "02_Expected_Deliverables_and_Outcomes.pdf")
    print("  Answer Key: 02_Expected_Deliverables_and_Outcomes")


def build_solution_paths(docx_to_pdf):
    doc = _newdoc(
        "Strategic Solution Paths",
        "Defensible recommendations and marking red lines",
    )
    add_body(
        doc,
        "There is no single correct strategy. Reward evidence-based reasoning. The paths "
        "below are strong; weak teams pick a slogan and ignore contradicting evidence.",
    )
    paths = [
        (
            "Path A - Depth before scale",
            "Pause expansion; fix Dong Nai and employer quality; negotiate a smaller quality-focused VPBank grant.",
            "Honest about D-02; addresses EP-02, ST-05, BN-02; sustainable operations.",
            "May leave funding on the table; must convince VPBank with a credible theory of change (DN-01).",
            "The ~1,100 on the waitlist (esp. Binh Duong / BN-04, YouthWorks leakage); Long An ambition and the "
            "Board growth faction; possibly part of the VPBank envelope if the grant is renegotiated down.",
        ),
        (
            "Path B - Conditional scale (most common strong answer)",
            "Accept VPBank but gate expansion behind quality milestones; fix Dong Nai first, then Long An.",
            "Balances donor pressure and field reality; stages risk; keeps the grant.",
            "Execution-heavy; requires M&E discipline teams must propose credibly.",
            "Whoever sits behind the gates - waitlisted youth until milestones clear; Long An if gates stay closed; "
            "young people who need help now while process/M&E work runs. Soft risk: staff burn if 'fix in parallel' "
            "becomes two full jobs.",
        ),
        (
            "Path C - Employer-led pivot",
            "Audit and certify employers; place only with vetted partners; make quality the brand.",
            "Answers EP-02 and ST-05 directly; differentiates from YouthWorks on outcomes not speed.",
            "Slower growth; depends on employer cooperation.",
            "Youth who would have taken weaker/exploitative placements for a short-term job; hubs that depend on "
            "volume employers (Dong Nai under pressure); anyone for whom fewer seats means a longer wait. Employers "
            "who fail the audit - and the young people those firms would have absorbed.",
        ),
        (
            "Path D - Digital hybrid",
            "Grow Pathway Digital for reach; keep Skills Forward for depth and placement quality.",
            "Lower unit cost (D-12); addresses waitlist pressure.",
            "Digital placement weaker; equity risk (BN-01); only strong with explicit caveats.",
            "Youth without devices/connectivity (BN-01); people who need in-person support (disability pathway, "
            "weaker digital skills - BN-03 risk); anyone steered into Digital who would have done better on Skills "
            "Forward's stronger placement quality. Waitlist may shrink on paper while equity quietly worsens.",
        ),
    ]
    for name, thesis, strengths, risks, left_behind in paths:
        add_heading(doc, name, 2)
        add_body(doc, "Thesis: " + thesis)
        add_body(doc, "Strengths: " + strengths)
        add_body(doc, "Who is left behind: " + left_behind, bold=True)
        add_body(doc, "Risks to probe in Q&A: " + risks)

    add_body(
        doc,
        "Across all four paths (if teams do not name it): young people with disability, Dong Nai participants "
        "stuck in bad internships, and anyone cherry-picked out to protect the 65% remain the red-line "
        "'left behind' groups - not a free pass for any path.",
        italic=True,
    )

    add_heading(doc, "Red lines - mark down if recommended", 1)
    for line in [
        "Reporting a shorter window than ninety days as the primary placement metric "
        "(contradicts D-03 Clauses 3-4 and DN-01; often floated as 'earlier visibility').",
        "Replacing paid caseworkers with volunteers (contradicts VR-01).",
        "National or provincial scale with no plan to fix Dong Nai.",
        "Cherry-picking easier-to-place youth to hit 65% without acknowledging the cost.",
        "Treating D-08 board emails as verified fact.",
        "Claiming attendance causes placement without naming confounders.",
        "Workflow A citing operational numbers without Request / Findings IDs.",
        "Workflow B sharing the raw workbook with A.",
    ]:
        add_bullet(doc, line)

    path = KEY / "Word" / "03_Strategic_Solution_Paths.docx"
    doc.save(path)
    docx_to_pdf(path, KEY / "03_Strategic_Solution_Paths.pdf")
    print("  Answer Key: 03_Strategic_Solution_Paths")


def build_inconsistencies_excel():
    wb = Workbook()

    ws = wb.active
    ws.title = "Dataset Issues"
    ws["A1"] = "GBF Case Study - Data Inconsistencies Register"
    ws["A1"].font = Font(name="Calibri", bold=True, size=14, color=NAVY)
    headers = ["Ref", "Issue", "Table", "Description", "Expected handling", "Found by team?"]
    for i, h in enumerate(headers, 1):
        c = ws.cell(row=3, column=i, value=h)
        c.font = Font(bold=True, color=WHITE)
        c.fill = PatternFill("solid", fgColor=TEAL)
        c.alignment = Alignment(wrap_text=True, vertical="center")
    for ri, row in enumerate(DATASET_ISSUES, 4):
        for ci, val in enumerate(row, 1):
            ws.cell(row=ri, column=ci, value=val)
        ws.cell(row=ri, column=6, value="")
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 42
    ws.column_dimensions["E"].width = 36
    ws.column_dimensions["F"].width = 14

    ws2 = wb.create_sheet("Document Contradictions")
    ws2["A1"] = "Document and narrative contradictions"
    ws2["A1"].font = Font(name="Calibri", bold=True, size=14, color=NAVY)
    headers2 = ["Ref", "Topic", "Sources", "Contradiction", "Strong team response"]
    for i, h in enumerate(headers2, 1):
        c = ws2.cell(row=3, column=i, value=h)
        c.font = Font(bold=True, color=WHITE)
        c.fill = PatternFill("solid", fgColor=TEAL)
    for ri, row in enumerate(DOCUMENT_ISSUES, 4):
        for ci, val in enumerate(row, 1):
            ws2.cell(row=ri, column=ci, value=val)
    ws2.column_dimensions["A"].width = 8
    ws2.column_dimensions["B"].width = 18
    ws2.column_dimensions["C"].width = 16
    ws2.column_dimensions["D"].width = 44
    ws2.column_dimensions["E"].width = 36

    ws3 = wb.create_sheet("Deliverable Standards")
    ws3["A1"] = "Expected deliverables and outcomes"
    ws3["A1"].font = Font(name="Calibri", bold=True, size=14, color=NAVY)
    headers3 = ["Deliverable", "Workflow", "Due", "Format", "Requirement", "Excellent", "Mark down if"]
    for i, h in enumerate(headers3, 1):
        c = ws3.cell(row=3, column=i, value=h)
        c.font = Font(bold=True, color=WHITE)
        c.fill = PatternFill("solid", fgColor=TEAL)
        c.alignment = Alignment(wrap_text=True, vertical="center")
    for ri, row in enumerate(DELIVERABLES, 4):
        for ci, val in enumerate(row, 1):
            cell = ws3.cell(row=ri, column=ci, value=val)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    for col, w in zip("ABCDEFG", [24, 10, 10, 10, 36, 36, 28]):
        ws3.column_dimensions[col].width = w

    out = KEY / "Data_Inconsistencies_Register.xlsx"
    wb.save(out)
    print(f"  Answer Key: {out.name}")


def build_all(docx_to_pdf):
    KEY.mkdir(parents=True, exist_ok=True)
    word_dir = KEY / "Word"
    if word_dir.exists():
        import shutil
        shutil.rmtree(word_dir)
    build_index(docx_to_pdf)
    build_inconsistencies_register(docx_to_pdf)
    build_expected_deliverables(docx_to_pdf)
    build_solution_paths(docx_to_pdf)
    build_inconsistencies_excel()
