#!/usr/bin/env python3
"""
Realistic document content for the Momentum / GBF case study.

Each builder writes a .docx into the staging folder and is registered in
MANIFEST with its final destination folder and output format:
  fmt = "pdf"  -> read-only reference material (converted to PDF)
  fmt = "docx" -> fill-in template (delivered as Word)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

from brand import (
    set_doc_defaults, add_cover, add_heading, add_body, add_bullet, add_number,
    add_table, add_meta_table, add_quote, add_signature_block, add_online_confirm,
    add_page_break,
)

ROOT = Path(__file__).resolve().parents[1]
STAGING = ROOT / "tools" / ".cache" / "word"
STAGING.mkdir(parents=True, exist_ok=True)

MANIFEST = []  # list of dicts: {name, dest, fmt}


def _doc(title, subtitle="", classification="Participant material", ref="", issued=""):
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc, title, subtitle, classification, ref, issued)
    return doc


def _save(doc, name, dest, fmt):
    path = STAGING / f"{name}.docx"
    doc.save(path)
    MANIFEST.append({"name": name, "dest": dest, "fmt": fmt})
    print(f"  built: {name} -> {dest} ({fmt})")
    return path


# ===========================================================================
# 0. START HERE
# ===========================================================================

def build_programme_handbook():
    doc = _doc(
        "Programme Guideline",
        "Handbook, weekly deliverables, roles, submission rules and ground rules in one place",
        ref="MOM-2026-GL-01", issued="Sunday 26 July 2026",
    )

    add_heading(doc, "1. Welcome to the engagement", 1)
    add_body(doc, "Your team has been retained as external strategy consultants by the Board "
             "of Generation Bridge Foundation (GBF), a Vietnamese non-profit working on youth "
             "unemployment. The engagement runs from Sunday 26 July to Sunday 20 September 2026 "
             "(eight working weeks, Monday to Friday). You will diagnose a genuine strategic "
             "problem, test your thinking against incomplete evidence, and present a "
             "recommendation to the Board.")
    add_body(doc, "This is a simulated but realistic engagement. There is no single correct "
             "answer. You will be assessed on the quality of your thinking, the honesty of your "
             "evidence, the discipline of your teamwork, and the clarity of your communication.")
    add_body(doc, "Generation Bridge Foundation, VPBank Foundation (as portrayed here), named "
             "staff, beneficiaries and employers are fictional constructs created for the "
             "Momentum Impact Case. They do not represent real organisations or people. Do not "
             "import outside data, news or personal knowledge into deliverables or inferences.",
             bold=True)

    add_heading(doc, "2. The client at a glance", 1)
    add_table(doc, ["Item", "Detail"], [
        ("Organisation", "Generation Bridge Foundation (Cau Noi The He) - fictional case client"),
        ("Founded", "2014, Ho Chi Minh City"),
        ("Mission", "Connect 16-24 year-olds who are not in education, employment or training (NEET) to dignified work"),
        ("2025 budget", "VND 18.2 billion"),
        ("Young people served in 2025", "800"),
        ("Current waitlist (Jan 2026)", "1,100"),
        ("Staff", "67 full-time, 12 part-time, ~340 active volunteers"),
        ("Delivery hubs", "Ho Chi Minh City (2), Binh Duong (1), Dong Nai (1), Long An (pilot)"),
    ], widths=[2.4, 4.2])

    add_heading(doc, "3. Why you have been retained", 1)
    add_body(doc, "The Board meets on Thursday 10 September 2026 to set GBF's direction for the "
             "next three years. Members do not agree. A major funder decision is on the table. "
             "Field staff and young people are describing problems that do not appear cleanly in "
             "the public materials. You have been retained to prepare a recommendation the Board "
             "can adopt - with options, evidence, and an honest view of trade-offs.")
    add_quote(doc, "We need a decision we can live with, not a slide that papers over our "
              "disagreements. Tell us what you would do in our place, and why.",
              "Engagement letter from the GBF Board, 15 January 2026")
    add_body(doc, "The engagement letter does not hand you a finished problem statement. Framing "
             "what the Board is really deciding - and what is noise - is part of the work "
             "(deliverable A1). Read the client data room carefully: senior staff, the largest "
             "donor and individual board members pull in different directions. Some of what you "
             "are told will not survive contact with the operational record held by Workflow B.")

    add_heading(doc, "4. How your team is organised: two workflows", 1)
    add_body(doc, "Every team of five or six splits into two consulting workflows. The split is "
             "deliberate. Each workflow holds incomplete information. You must talk to each other "
             "to converge. Read both engagement briefs.")
    add_table(doc, ["Workflow", "Owns", "Folder"], [
        ("A - Impact Strategy", "Problem framing, stakeholder synthesis, the questions worth asking, "
         "strategic options, and the Board narrative",
         "2_Workflow_A_Impact_Strategy"),
        ("B - Operations & Analytics", "The operational datasets, data dictionary, cleaning, "
         "analysis, and plain-language translation of findings",
         "3_Workflow_B_Operations_and_Analytics"),
    ], widths=[1.7, 3.0, 2.0])
    add_heading(doc, "Suggested roles (assign in Week 1)", 2)
    add_table(doc, ["Role", "Count", "Owns"], [
        ("Engagement Lead", "1", "Timeline, board narrative, final integration, Request Log health"),
        ("Workflow A Lead", "1", "Issue tree, options, trade-off reflection, Analysis Requests"),
        ("Research Lead", "1", "Interview synthesis, credibility matrix, validation"),
        ("Workflow B Lead", "1", "Data dictionary, cleaning, Findings Memos, dashboard"),
        ("Operations Analyst", "1", "Funding scenarios, cost analysis"),
        ("Project & Quality Lead", "1", "Minutes, version control, handoff checklist, file naming"),
    ], widths=[1.9, 0.7, 3.8])
    add_body(doc, "In a team of five, combine Research Lead with Workflow A Lead. Full ways of "
             "working (decision rules, free-riding prevention) are also in "
             "4_Shared_Toolkit/Team_Roles_and_Ways_of_Working.pdf.")

    add_heading(doc, "The A↔B evidence loop", 2)
    add_body(doc, "Workflow A does not receive the operational datasets and must not ask "
             "facilitators for the file. Workflow B does not push raw tables into A's narrative. "
             "Instead the team runs a four-step loop (templates in 4_Shared_Toolkit/Templates/):")
    add_table(doc, ["Step", "Who", "What", "Template / output"], [
        ("1. Data Estate Brief", "B → A",
         "Catalog of what B holds (tables, grain, coverage, limits in plain language). "
         "Not the workbook, not the data dictionary, not field dumps that let A write the analysis.",
         "Data_Estate_Brief.docx → PDF in A_B_Exchange/"),
        ("2. Demand Brief", "A → B",
         "Ranked decision needs and hypotheses (what would change the recommendation). "
         "Demand, not a shopping list of column names.",
         "Demand_Brief.docx → PDF in A_B_Exchange/"),
        ("3. Analysis Plan", "B (share summary with A)",
         "B owns how to answer each priority demand: approach, definitions, risks, ETA. "
         "A does not co-edit pivots or receive interim raw extracts.",
         "Analysis_Plan.docx → PDF in A_B_Exchange/"),
        ("4. Request → Findings", "A then B",
         "Formal Analysis Request (R-xx) and Findings Memo (FM-xx) for each answered demand. "
         "This is the assessment trail markers read.",
         "Analysis_Request.docx / Findings_Memo.docx"),
    ], widths=[1.3, 1.0, 2.4, 1.7])
    add_body(doc, "Before final options, the team must complete at least three Analysis Requests "
             "with Findings Memos returned. Every quantitative claim in A's deliverables must "
             "cite a Request ID or Findings Memo. The Data Estate Brief, Demand Brief and "
             "Analysis Plan must exist in A_B_Exchange/ before R-01 is treated as complete. "
             "Before any gated submission (options, final report, board pack), complete the "
             "Cross-Workflow Handoff Checklist - see Section 8.",
             bold=True)

    add_heading(doc, "5. What you will produce - by week", 1)
    add_body(doc, "Programme opens Sunday 26 July 2026 and closes Sunday 20 September 2026. "
             "Working weeks run Monday to Friday. Weekly submissions are due Friday 17:00 ICT "
             "unless a Board or handoff time is stated. Day-by-day detail lives only in "
             "GBF_Consulting_Toolkit.xlsx (Master Timeline).")
    add_table(doc, ["Week", "Dates (2026)", "Submit by end of week"], [
        ("1", "27-31 July",
         "Kickoff; roles confirmed; Project Charter & RACI; Request Log set up; start data-room reading"),
        ("2", "3-7 August",
         "A1 issue tree; A1b stakeholder map; information credibility matrix; B profiling under way "
         "(Data Estate Brief draft)"),
        ("3", "10-14 August",
         "B Data Estate Brief (Mon); A Demand Brief (Tue after reading DEB); B Analysis Plan; "
         "R-01 filed + FM-01 returned; "
         "B1 data quality (B internal); A2 interview synthesis; A3 empathy & journey maps; B2 dashboard v1"),
        ("4", "17-21 August",
         "R-02 and R-03 + FM-02/FM-03 (each with Analysis Plan updates as needed); B3 analysis summary; "
         "A4 strategic options (cite R/FM IDs); B4 evidence-check memo"),
        ("5", "24-28 August",
         "A5 trade-off reflection; A6 validation; B5 funding scenarios; A7 revision log; "
         "executive summary draft"),
        ("6", "31 Aug-4 Sep",
         "Final report draft for handoff; online handoff checklist confirmed"),
        ("7", "7-11 September",
         "BOARD PRESENTATION Thu 10 Sep 09:00; all final deliverables Fri 11 Sep 17:00"),
        ("8", "14-18 September",
         "Peer evaluation (Mon); reflection journal (Tue); lessons-learned; programme closes Sun 20 Sep"),
    ], widths=[0.7, 1.5, 4.2])
    add_body(doc, "Track the loop on the Request Log tab: R/FM rows plus milestone ticks for "
             "Data Estate Brief, Demand Brief and first Analysis Plan. "
             "Minimum three completed R/FM pairs before final options.")

    add_heading(doc, "6. How you will be assessed", 1)
    add_body(doc, "A 100-point rubric applies. The table below is the expected scale for planning "
             "your effort. Facilitators may refine descriptors or band cut-offs; any change will "
             "be announced. Treat these weights as indicative, not a frozen contract.",
             italic=True)
    add_table(doc, ["Criterion", "Points (expected)"], [
        ("Problem understanding", "15"), ("Research quality", "10"),
        ("Design thinking", "15"), ("Data analysis and translation", "20"),
        ("Recommendations", "20"), ("Presentation", "10"),
        ("Project management", "5"), ("Teamwork and cross-workflow exchange", "5"),
    ], widths=[3.6, 1.8])
    add_body(doc, "Presentation Day also applies a panel split (Buddy / Core Team / Guest Board). "
             "Your Buddy will confirm the live weights at kickoff. Strong teams ask precise "
             "questions, translate findings honestly, and own the costs of their preferred "
             "recommendation.")

    add_heading(doc, "7. Citation standard", 1)
    add_body(doc, "Every claim must be trackable for markers. Do not cite a document code alone.")
    add_table(doc, ["Source type", "Required citation form", "Example"], [
        ("Client document", "Ref + section or page heading", "D-02 §1 Dong Nai; D-01 'Our year in numbers'"),
        ("Interview transcript", "Transcript code + speaker", "BN-02 (Tuan); ST-05 (Minh)"),
        ("Findings / requests", "Request or Findings ID", "R-02; FM-03"),
        ("Dataset (B only, via memo)", "Table + field + definition used", "employment_outcomes.placed_90d (formal)"),
    ], widths=[1.6, 2.4, 2.4])

    add_heading(doc, "8. How you receive the package, work, and submit", 1)
    add_heading(doc, "Receiving the package", 2)
    add_body(doc, "Your Buddy (or programme Core Team) issues the Participants/ folder to your "
             "team at kickoff - usually as a Drive share or a ZIP. You do not need to download "
             "anything from a public repository. If your folder is incomplete, ask your Buddy "
             "before inventing missing files.")
    add_heading(doc, "Shared workspace", 2)
    add_body(doc, "Each team works in the Google Drive / Google Workspace folder created by the "
             "programme (or Microsoft 365 if your cohort is assigned that stack). Buddies and "
             "Core Team only review files in that folder and via the weekly submission Form. "
             "Copy templates from the package into your Drive before editing - keep the issued "
             "package as a read-only reference.")
    add_heading(doc, "File naming (mandatory)", 2)
    add_body(doc, "Use: TeamShortName_Owner_DeliverableCode_Version.ext")
    add_table(doc, ["Example", "Meaning"], [
        ("Tiger_All_Charter_v1.pdf", "Whole-team Project Charter, version 1"),
        ("Tiger_WSA_A1_IssueTree_v2.pdf", "Workflow A deliverable A1"),
        ("Tiger_WSB_DEB_v1.pdf", "Data Estate Brief (B → A)"),
        ("Tiger_WSA_DMD_v1.pdf", "Demand Brief (A → B)"),
        ("Tiger_WSB_AP01_Plan_v1.pdf", "Analysis Plan AP-01 (B)"),
        ("Tiger_WSA_R01_Request_v1.pdf", "Analysis Request R-01 (filed by A)"),
        ("Tiger_WSB_FM01_Findings_v1.pdf", "Findings Memo FM-01 (returned by B)"),
        ("Tiger_All_Handoff_v1.pdf", "Completed Cross-Workflow Handoff Checklist"),
        ("Tiger_All_BoardDeck_vFinal.pptx", "Final board deck"),
    ], widths=[3.2, 3.2])
    add_heading(doc, "A_B_Exchange folder (mandatory for assessment)", 2)
    add_body(doc, "Create a folder in your team Drive named A_B_Exchange. Every official "
             "cross-workflow communication must be saved there as PDF using the templates below. "
             "Markers and Buddies assess teamwork from this folder - not from chat history.")
    add_table(doc, ["What", "Template location", "Who files"], [
        ("Data Estate Brief (DEB)", "4_Shared_Toolkit/Templates/Data_Estate_Brief.docx", "Workflow B"),
        ("Demand Brief (DMD)", "4_Shared_Toolkit/Templates/Demand_Brief.docx", "Workflow A"),
        ("Analysis Plan (AP-xx)", "4_Shared_Toolkit/Templates/Analysis_Plan.docx", "Workflow B"),
        ("Analysis Request (R-xx)", "4_Shared_Toolkit/Templates/Analysis_Request.docx", "Workflow A"),
        ("Findings Memo (FM-xx)", "4_Shared_Toolkit/Templates/Findings_Memo.docx", "Workflow B"),
    ], widths=[2.0, 3.2, 1.4])
    add_number(doc, "Export each completed Word template to PDF before filing.")
    add_number(doc, "Name files with DEB / DMD / AP-xx / R-xx / FM-xx codes.")
    add_number(doc, "Update the Request Log tab in GBF_Consulting_Toolkit.xlsx whenever a file is added.")
    add_number(doc, "Informal chat is fine for logistics only. Decisions, asks and answers that "
               "affect the recommendation must appear in A_B_Exchange.")
    add_number(doc, "Sequence: Data Estate Brief → Demand Brief → Analysis Plan → R-xx → FM-xx. "
               "Minimum three completed R/FM pairs before final options.")

    add_heading(doc, "Cross-Workflow Handoff Checklist (gate before submission)", 2)
    add_body(doc, "Where to find it: 4_Shared_Toolkit/Templates/Cross_Workflow_Handoff_Checklist.docx "
             "(also listed in the Submission Checklist tab of GBF_Consulting_Toolkit.xlsx).")
    add_body(doc, "When it is required: before submitting strategic options (A4), the final report "
             "draft (Week 6), and the final board pack (Week 7). Master Timeline task W6-02 is "
             "the formal online confirmation checkpoint (Wed 2 Sep, 12:00 ICT) - do not leave it "
             "to the last hour before the Board.")
    add_body(doc, "How online confirmation works (no wet signature, no printing):")
    add_number(doc, "Copy the Word template into your team Drive and fill the checklist boxes "
               "([x] for done).")
    add_number(doc, "Section A: a Workflow A reviewer confirms Workflow B's Findings quality.")
    add_number(doc, "Section B: a Workflow B reviewer confirms Workflow A's evidence cites and options.")
    add_number(doc, "Section C: the Engagement Lead confirms only if A_B_Exchange/ and the Request "
               "Log show at least three genuine R/FM pairs (not a last-minute dump).")
    add_number(doc, "Each confirmer completes the online block at the end of their section: tick "
               "[x], type full name, type date (YYYY-MM-DD), and note the channel "
               "(e.g. 'Drive file' or 'comment on this Doc').")
    add_number(doc, "Export the completed checklist to PDF, name it Team_All_Handoff_vN.pdf, save "
               "it in your team Drive root (or a /Handoff/ folder), and submit the Drive link via "
               "the weekly Google Form when the timeline requires it.")
    add_number(doc, "Your Buddy checks that the PDF exists and the three confirmation blocks are "
               "filled - they do not re-do your quality review.")

    add_heading(doc, "Submission channel", 2)
    add_number(doc, "Upload the named file to your team Drive folder.")
    add_number(doc, "Submit the Drive link (or attach if under size limit) via the programme "
               "Google Form for that week - your Buddy will share the Form URL at kickoff.")
    add_number(doc, "Email is backup only if the Form is down; always CC your Buddy.")
    add_number(doc, "When submitting options or the final report, include the link to A_B_Exchange "
               "and to the completed Handoff Checklist PDF so markers can see the trail.")

    add_heading(doc, "9. Ground rules", 1)
    add_number(doc, "Fictional case only: do not treat GBF or named individuals as real; do not "
               "bring outside data, news or personal research into deliverables or inferences.")
    add_number(doc, "Use only the materials in this package.")
    add_number(doc, "Workflow A must not open or request the raw operational datasets. Route all "
               "quantitative needs through Analysis Requests to Workflow B.")
    add_number(doc, "Workflow B must not share the raw workbook with Workflow A. Share Findings "
               "Memos and curated charts only.")
    add_number(doc, "Official A-B evidence lives only in A_B_Exchange/ using the Data Estate Brief, "
               "Demand Brief, Analysis Plan, Request and Findings Memo templates. Chat is not "
               "the assessment record.")
    add_number(doc, "Cite every claim to the citation standard in Section 7.")
    add_number(doc, "State the definitions you use, especially for the placement rate.")
    add_number(doc, "Fill templates in Word using Nunito or Lexend (Vietnamese-safe). Present in "
               "PowerPoint with the same fonts. Export read-only finals as PDF.")
    add_number(doc, "Handoff: complete 4_Shared_Toolkit/Templates/Cross_Workflow_Handoff_Checklist.docx "
               "with online ticks (typed name + date) - see Guideline Section 8. No wet signature.")
    add_number(doc, "Your Buddy is a facilitator, not a consultant. They will not give you the "
               "answer or the data file.")

    add_heading(doc, "10. Package contents", 1)
    add_table(doc, ["Folder", "What is inside"], [
        ("0_Start_Here", "This Programme Guideline; data room index"),
        ("1_Client_Data_Room", "GBF documents D-01, D-02, D-03, D-08, D-10, D-11, D-12 and transcripts (no datasets)"),
        ("2_Workflow_A_Impact_Strategy", "Workflow A brief, playbook, A templates"),
        ("3_Workflow_B_Operations_and_Analytics", "Workflow B brief, data dictionary, datasets, dashboard, B templates"),
        ("4_Shared_Toolkit", "Toolkit workbook (Request Log + Submission Checklist), guides, board deck, shared templates including Handoff Checklist"),
    ], widths=[2.9, 3.6])

    _save(doc, "00_Programme_Guideline", "0_Start_Here", "pdf")


def build_data_room_index():
    doc = _doc("Data Room Index", "A guide to documents and transcripts in the shared client room",
               ref="MOM-2026-DRI-01", issued="Sunday 26 July 2026")
    add_body(doc, "This index lists what is in the shared client data room. In a real engagement, "
             "clients hand you mixed-quality material and expect you to decide what to trust. "
             "We do not pre-rate reliability for you. That is your job in the Information "
             "Credibility Matrix (Week 2). Operational datasets are not in this folder; they "
             "are held by Workflow B under the A↔B evidence loop (Estate → Demand → Plan → R/FM).")

    add_heading(doc, "How to question what you read", 1)
    add_body(doc, "Before you lean on any claim, ask:")
    add_number(doc, "Who produced this, for whom, and with what incentive to look good or look candid?")
    add_number(doc, "Are headline metrics defined? Would another definition change the story?")
    add_number(doc, "Does another document or interview say something different about the same topic?")
    add_number(doc, "Is this a measured fact, a management opinion, or a stakeholder preference?")
    add_number(doc, "What would you need Workflow B to check in the operational record before you "
               "treat a number as settled?")
    add_body(doc, "Capture your judgements in the credibility matrix. Public reports, internal "
             "memos, emails, surveys, audited accounts and interviews are not automatically "
             "equal. Disagreement between sources is normal; resolving it is part of the work.")

    add_heading(doc, "How to cite these sources", 1)
    add_body(doc, "Markers will not hunt through an entire PDF for your claim. Cite document "
             "reference plus section heading or page topic, not the code alone. Transcripts: cite "
             "code and speaker. Examples: (D-02, §1 Dong Nai); (D-01, 'Our year in numbers'); "
             "(BN-02, Tuan).")

    add_heading(doc, "Client documents", 1)
    add_table(doc, ["Ref", "Document", "Type / audience"], [
        ("D-01", "Annual Report 2024", "Public report (donors, partners, public)"),
        ("D-02", "Programme Review 2025", "Internal memo (senior staff)"),
        ("D-03", "VPBank Foundation grant agreement (excerpt)", "Funder contract excerpt"),
        ("D-08", "Board email thread", "Informal Board correspondence"),
        ("D-10", "Competitor brief: YouthWorks", "Internal competitive note"),
        ("D-11", "Staff engagement survey 2025", "Anonymous staff survey summary"),
        ("D-12", "Financial statements 2023-2025", "Audited accounts"),
    ], widths=[0.6, 3.2, 2.8])

    add_heading(doc, "Interview transcripts", 1)
    add_body(doc, "The full transcripts are in '20_Interview_Transcripts'. Fourteen "
             "conversations were recorded in January 2026 by the engagement team and lightly "
             "edited for clarity (not stripped of natural speech). Cite as CODE (Speaker), "
             "e.g. ST-01 (Hung). Treat each voice as one perspective, not organisational truth.")
    add_table(doc, ["Ref", "Interviewee", "Role"], [
        ("ST-01", "Hung Vo", "Programmes Director"),
        ("ST-02", "Thao Le", "Operations Director"),
        ("ST-03", "Linh Nguyen", "Monitoring & Evaluation Officer"),
        ("ST-04", "Anh Duc Pham", "Development Director"),
        ("ST-05", "Minh Tran", "Field Officer, Dong Nai hub"),
        ("BN-01", "Mai", "Programme graduate, Ho Chi Minh City"),
        ("BN-02", "Tuan", "Withdrew from programme, Dong Nai"),
        ("BN-03", "Lan", "Graduate with a disability, Ho Chi Minh City"),
        ("BN-04", "Hang", "On the waitlist, Binh Duong"),
        ("DN-01", "Ms. Huong", "Programme Officer, VPBank Foundation"),
        ("EP-01", "Mr. Khang", "HR Manager, gold-tier employer partner"),
        ("EP-02", "Ms. Trang", "Ops manager, logistics subcontractor (anonymous)"),
        ("VR-01", "Chi Hoa", "Volunteer mentor, three years' service"),
        ("BD-01", "Mr. T.N.", "Member of the GBF Board"),
    ], widths=[0.7, 2.0, 3.9])

    add_heading(doc, "Operational datasets (Workflow B only)", 1)
    add_body(doc, "The monitoring export and data dictionary sit in Workflow B's folder. "
             "Workflow A does not receive these files. If you need a number, a comparison, or a "
             "test of a hypothesis, file an Analysis Request. Figures that appear in shared "
             "documents are starting points for questions, not a substitute for querying the "
             "operational record.")

    _save(doc, "01_Data_Room_Index", "0_Start_Here", "pdf")


def build_deadlines():
    """Deprecated: weekly deadlines live in 00_Programme_Guideline. Kept as no-op for safety."""
    return


# ===========================================================================
# 1. CLIENT DATA ROOM
# ===========================================================================

def build_wsa_start():
    doc = _doc(
        "Workflow A - Start Here",
        "Impact Strategy: what is in this folder and what you must deliver",
        classification="Workflow A",
        ref="MOM-2026-WSA-00",
        issued="Sunday 26 July 2026",
    )
    add_body(doc, "You are on Workflow A (Impact Strategy). Read this page first, then the "
             "engagement brief and design-thinking playbook. You set direction and frame the "
             "questions. You do not have access to the operational datasets. Do not ask "
             "facilitators for the file. Read B's Data Estate Brief, file a Demand Brief, then "
             "route quantitative needs through Analysis Requests. B owns the Analysis Plan.")
    add_heading(doc, "Files in this folder", 1)
    add_table(doc, ["File", "Purpose"], [
        ("WSA_Engagement_Brief.pdf", "Your mandate, deliverables and constraints"),
        ("WSA_Design_Thinking_Playbook.pdf", "Run sheets for workshops and activities"),
        ("Templates/", "Word files to complete and export to PDF before submission"),
    ], widths=[2.8, 3.6])
    add_body(doc, "Data Estate Brief, Demand Brief, Analysis Plan, Analysis Request and Findings "
             "Memo templates live in 4_Shared_Toolkit/Templates/. Log every step on the Request "
             "Log. Save every PDF in A_B_Exchange/ - that folder is what markers assess.")
    add_heading(doc, "Your deliverables", 1)
    add_table(doc, ["Ref", "Deliverable", "Template", "Due (17:00 ICT)"], [
        ("A1", "Problem statement and issue tree", "Problem_Statement_and_Issue_Tree.docx", "Wed 5 Aug"),
        ("A1b", "Stakeholder map", "Stakeholder_Map.docx", "Thu 6 Aug"),
        ("C", "Information credibility matrix (shared)", "4_Shared_Toolkit/Templates/Information_Credibility_Matrix.docx", "Thu 6 Aug"),
        ("DMD", "Demand Brief (to B)", "4_Shared_Toolkit/Templates/Demand_Brief.docx", "Tue Week 3"),
        ("A2", "Interview synthesis", "Interview_Synthesis.docx", "Wed 12 Aug"),
        ("A3", "Empathy and journey maps", "Empathy_and_Journey_Map.docx", "Wed 12 Aug"),
        ("A4", "Strategic options and recommendation", "Strategic_Options.docx", "Thu 20 Aug"),
        ("A5", "Trade-off reflection", "Tradeoff_Reflection.docx", "Wed 26 Aug"),
        ("A6", "Validation protocol", "Validation_Protocol.docx", "Wed 26 Aug"),
        ("A7", "Recommendation revision log", "Recommendation_Revision_Log.docx", "Thu 27 Aug"),
    ], widths=[0.5, 2.4, 2.6, 1.1])
    add_heading(doc, "Before gated submissions (options, finals, board pack)", 1)
    add_number(doc, "Every insight from transcripts cites CODE (Speaker), e.g. BN-02 (Tuan).")
    add_number(doc, "Every document cite includes section/heading, e.g. D-02 §1 Dong Nai.")
    add_number(doc, "Every number cites a Findings Memo / Request ID from Workflow B.")
    add_number(doc, "Data Estate Brief, Demand Brief and Analysis Plan PDFs are in A_B_Exchange/ "
               "(required before R-01 is treated as complete).")
    add_number(doc, "At least three Analysis Requests have Findings Memos returned "
               "(PDFs in A_B_Exchange/) before final options.")
    add_number(doc, "For options / final report / board pack: Workflow B has completed Section B of "
               "4_Shared_Toolkit/Templates/Cross_Workflow_Handoff_Checklist.docx "
               "(online tick + typed name/date - see Guideline Section 8; no wet signature).")
    _save(doc, "WSA_Start_Here", "2_Workflow_A_Impact_Strategy", "pdf")


def build_wsb_start():
    doc = _doc(
        "Workflow B - Start Here",
        "Operations and Analytics: what is in this folder and what you must deliver",
        classification="Workflow B",
        ref="MOM-2026-WSB-00",
        issued="Sunday 26 July 2026",
    )
    add_body(doc, "You are on Workflow B (Operations and Analytics). You alone hold the "
             "operational datasets and the data dictionary. You also have full access to the "
             "shared client data room (1_Client_Data_Room/) and the Programme Guideline - read "
             "those like any consultant would. Your job is to clean, analyse, brief A with a "
             "Data Estate Brief (catalog only), own Analysis Plans against A's Demand, and "
             "return Findings Memos in plain language. Do not share the raw workbook with A.")
    add_heading(doc, "What you can open", 1)
    add_table(doc, ["Location", "Access"], [
        ("0_Start_Here/", "Yes - Guideline and Data Room Index"),
        ("1_Client_Data_Room/", "Yes - D-01/02/03/08/10/11/12 and all interview transcripts"),
        ("3_Workflow_B_.../ (this folder)", "Yes - datasets, dictionary, dashboard, B templates"),
        ("4_Shared_Toolkit/", "Yes - Request Log, A_B_Exchange templates, board deck"),
        ("2_Workflow_A_.../ Templates", "Read if useful; do not own A's narrative deliverables"),
        ("GBF_Datasets.xlsx", "B only - never share the raw file with A"),
    ], widths=[3.2, 3.2])
    add_heading(doc, "Files in this folder", 1)
    add_table(doc, ["File", "Purpose"], [
        ("WSB_Engagement_Brief.pdf", "Your mandate, deliverables and constraints"),
        ("WSB_Data_Dictionary.pdf", "Schema and field notes for the export you hold"),
        ("WSB_Data_Analysis_Workbook.pdf", "Guidance for analysis tasks DB-1 to DB-5"),
        ("GBF_Datasets.xlsx", "Operational export (B only - do not share raw)"),
        ("GBF_Performance_Dashboard.xlsx", "Build your working dashboard here"),
        ("Templates/", "Word files to complete and export to PDF before submission"),
    ], widths=[2.8, 3.6])
    add_heading(doc, "Your deliverables", 1)
    add_table(doc, ["Ref", "Deliverable", "Where", "Due (17:00 ICT)"], [
        ("DEB", "Data Estate Brief (to A)", "4_Shared_Toolkit/Templates/Data_Estate_Brief.docx", "Mon Week 3"),
        ("B1", "Data quality report (internal)", "Templates/Data_Quality_Report.docx", "Tue 11 Aug"),
        ("AP", "Analysis Plans", "4_Shared_Toolkit/Templates/Analysis_Plan.docx", "From Week 3"),
        ("B2", "Performance dashboard (working file)", "GBF_Performance_Dashboard.xlsx", "Fri 14 Aug"),
        ("B3", "Analysis summary + Findings Memos", "Templates/Analysis_Summary.docx + Findings_Memo", "Wed 19 Aug"),
        ("B4", "Evidence-check memo on A's options", "Templates/Evidence_Check_Memo.docx", "Fri 21 Aug"),
        ("B5", "Funding scenarios", "Shared Toolkit / Funding Scenarios tab", "Thu 27 Aug"),
    ], widths=[0.5, 2.2, 2.4, 1.3])
    add_heading(doc, "Before gated submissions (options, finals, board pack)", 1)
    add_number(doc, "State the placement definition you used and why.")
    add_number(doc, "Document every cleaning decision in the Assumption Log.")
    add_number(doc, "Data Estate Brief is a catalog for A - never attach the workbook or dictionary.")
    add_number(doc, "Analysis Plans show how you will answer Demand / R-xx; you own the method.")
    add_number(doc, "Findings Memos use plain language; charts are curated, not a dump of pivots.")
    add_number(doc, "Every DEB / AP / FM PDF is saved in A_B_Exchange/ and logged on the Request Log.")
    add_number(doc, "Do not claim causation from correlation on attendance or mentoring.")
    add_number(doc, "For options / final report / board pack: Workflow A has completed Section A of "
               "4_Shared_Toolkit/Templates/Cross_Workflow_Handoff_Checklist.docx "
               "(online tick + typed name/date - see Guideline Section 8; no wet signature).")
    _save(doc, "WSB_Start_Here", "3_Workflow_B_Operations_and_Analytics", "pdf")


def build_d01_annual_report():
    doc = _doc("Annual Report 2024", "Generation Bridge Foundation (public edition, excerpt)",
               classification="Client document D-01", ref="D-01", issued="Published April 2025")

    add_meta_table(doc, [
        ("Organisation", "Generation Bridge Foundation (Cau Noi The He)"),
        ("Registered office", "47 Nguyen Van Troi, Phu Nhuan District, Ho Chi Minh City"),
        ("Registration", "Decision 124/QD-SLDTBXH, 2017"),
        ("Reporting period", "1 January to 31 December 2024"),
        ("Publication", "April 2025"),
    ])
    add_heading(doc, "A message from our Executive Director", 1)
    add_body(doc, "In a year of rising need, Generation Bridge Foundation stood beside more young "
             "people than ever before. Across four provinces, 592 young men and women completed "
             "one of our programmes, and I am proud to report that 71 per cent of them were in "
             "paid work within ninety days - comfortably ahead of our 65 per cent target.")
    add_body(doc, "None of this is possible without our partners. I want to thank the VPBank "
             "Foundation, whose multi-year commitment underpins our flagship Skills Forward "
             "programme, and the hundreds of volunteer mentors who give their evenings and "
             "weekends to walk alongside our young people.")
    add_body(doc, "Demand for our services continues to outstrip what we can offer. As I write, "
             "more than a thousand young people are waiting for a place. Meeting that need, "
             "without compromising the quality that makes our model work, is the central "
             "challenge of the years ahead.")
    add_quote(doc, "We do not simply train young people and hope. We connect them to real "
              "employers, real internships, and real jobs.", "Dr. Lan Phuong Nguyen, Executive Director")

    add_heading(doc, "Our year in numbers", 1)
    add_table(doc, ["Indicator", "2023", "2024"], [
        ("Young people completing a programme", "495", "592"),
        ("Job placement rate within 90 days", "68%", "71%"),
        ("Active employer partners", "48", "55"),
        ("Active volunteer mentors", "80", "95"),
        ("Provinces served", "3", "4"),
    ], widths=[3.2, 1.5, 1.5])
    add_body(doc, "Placement is defined as any paid employment of at least fifteen hours per "
             "week within ninety days of completing a programme. This includes flexible and "
             "platform-based work such as ride-hailing and online selling.", italic=True, size=10)

    add_heading(doc, "Our programmes", 1)
    add_body(doc, "Skills Forward, our flagship, combines eight weeks of employability training "
             "with an eight-week paid internship at a partner employer. Pathway Digital offers a "
             "twelve-week online course in e-commerce and digital administration, followed by a "
             "four-week in-person lab. Bridge Mentors pairs young people with trained volunteer "
             "mentors for six months of one-to-one support.")

    add_heading(doc, "Looking ahead", 1)
    add_body(doc, "In 2025 and beyond we intend to deepen our employer relationships, strengthen "
             "our monitoring, and explore how we can responsibly reach the thousand-plus young "
             "people still waiting for a place.")
    _save(doc, "D01_Annual_Report_2024", "1_Client_Data_Room", "pdf")


def build_d02_programme_review():
    doc = _doc("Programme Review 2025", "Internal memorandum - not for external circulation",
               classification="Client document D-02 (internal)", ref="D-02")
    add_meta_table(doc, [
        ("To", "Dr. Lan Phuong Nguyen, Executive Director"),
        ("From", "Hung Vo, Programmes Director"),
        ("Date", "8 December 2025"),
        ("Subject", "Honest review of 2025 delivery ahead of the Board strategy session"),
        ("Classification", "Internal - candid; please do not forward"),
    ])
    add_body(doc, "Lan, ahead of the January strategy session you asked me for a frank picture "
             "rather than the version we put in front of donors. Here it is.")

    add_heading(doc, "1. Dong Nai is in trouble", 1)
    add_body(doc, "The Dong Nai hub is our weakest performer and the gap is widening. Internship "
             "drop-out reached 38 per cent last quarter against an organisation-wide average of "
             "22 per cent. Two causes stand out. First, transport: young people are commuting 25 "
             "to 40 kilometres to industrial-zone placements and paying VND 150,000 to 200,000 a "
             "month in bus fares that we do not subsidise. Second, employer behaviour: several "
             "partners treat interns as free labour, with long hours and no clear duties. I have "
             "not shared the full Dong Nai numbers with VPBank.")

    add_heading(doc, "2. Pathway Digital is being reported too generously", 1)
    add_body(doc, "The completion rate we reported to the donor was 68 per cent. The true "
             "figure, measured as completers divided by those who actually started, is 52 per "
             "cent. The digital team used enrolments as the denominator. We need to correct this "
             "before it becomes a reputational problem.")

    add_heading(doc, "3. The Long An pilot has not worked", 1)
    add_body(doc, "Drop-out in the Long An pilot is around 60 per cent. The catchment is thinner "
             "than we assumed and we do not have enough local employers. I would not expand there "
             "until the model is fixed elsewhere.")

    add_heading(doc, "4. Our people are stretched", 1)
    add_body(doc, "Field staff turnover was 34 per cent this year. The team is doing heroic work, "
             "but it is not sustainable, and every departure costs us employer relationships that "
             "took years to build.")

    add_heading(doc, "Recommendation", 1)
    add_body(doc, "Before we promise anyone growth, we should fix Dong Nai, correct our "
             "reporting, and invest in the quality of our employer partnerships. Scaling a "
             "broken internship experience multiplies harm for the young people least able to "
             "absorb it. I know this is not the story Development wants to tell. It is the true one.")
    add_body(doc, "Hung")
    _save(doc, "D02_Programme_Review_Internal_Memo", "1_Client_Data_Room", "pdf")


def build_d03_grant_agreement():
    doc = _doc("Conditional Grant Agreement", "VPBank Foundation and Generation Bridge Foundation (excerpt)",
               classification="Client document D-03", ref="D-03", issued="Draft, dated 26 September 2025")
    add_body(doc, "This excerpt reproduces the key commercial terms of the conditional renewal "
             "offered by the VPBank Foundation. Legal boilerplate has been omitted.")
    add_meta_table(doc, [
        ("Grantor", "VPBank Foundation"),
        ("Grantee", "Generation Bridge Foundation"),
        ("Grant value", "VND 8,000,000,000 (eight billion dong)"),
        ("Term", "1 January 2026 to 31 December 2028 (three years)"),
        ("Disbursement", "Annual tranches, released on achievement of conditions"),
    ])

    add_heading(doc, "Clause 3 - Conditions of disbursement", 1)
    add_body(doc, "Each annual tranche is conditional on the Grantee demonstrating, to the "
             "reasonable satisfaction of the Grantor, that:")
    add_number(doc, "The job placement rate is at least 65 per cent, measured as the proportion "
               "of programme completers in paid employment of fifteen hours per week or more, "
               "ninety days after completion;")
    add_number(doc, "The Grantee is on a credible trajectory to serve a cumulative total of "
               "4,000 young people by 31 December 2028; and")
    add_number(doc, "The Grantee maintains a monitoring and evaluation framework acceptable to "
               "the Grantor, with quarterly reporting.")

    add_heading(doc, "Clause 4 - Measurement", 1)
    add_body(doc, "Placement for the purpose of Clause 3 shall be assessed ninety days after "
             "programme completion. The Grantor will accept flexible and platform-based "
             "employment where it is stable and provides at least fifteen hours of work per "
             "week, but requests that formal, contracted employment be reported separately. "
             "Management may track earlier leading indicators for internal use; such indicators "
             "do not replace the ninety-day measure for disbursement decisions.")

    add_heading(doc, "Clause 6 - Use of funds", 1)
    add_body(doc, "Funds are provided for the expansion of the Grantee's employment programmes, "
             "including the establishment of new delivery capacity. The Grantor supports, but "
             "does not require, expansion into new provinces, provided that quality standards are "
             "maintained.")
    add_body(doc, "Note appended by the GBF Development team: 'This is the largest single "
             "commitment in our history. We should accept.'", italic=True, size=10)
    _save(doc, "D03_VPBank_Grant_Agreement", "1_Client_Data_Room", "pdf")


def build_d08_email_thread():
    doc = _doc("Board Correspondence", "Email thread forwarded to the engagement team",
               classification="Client document D-08 (unverified)", ref="D-08")
    add_body(doc, "The following thread was forwarded to the engagement team by a member of "
             "staff who asked not to be named. It has not been verified. Treat it as an "
             "indication of Board sentiment, not as fact.", italic=True, size=10)

    def email(frm, to, date, subject, body_paras):
        add_meta_table(doc, [("From", frm), ("To", to), ("Date", date), ("Subject", subject)])
        for p in body_paras:
            add_body(doc, p)
        doc.add_paragraph()

    email("T.N. (Board member)", "GBF Board distribution list", "Mon 5 Jan 2026, 21:14",
          "We cannot afford to be timid",
          ["Colleagues, I will say what others are thinking. Lan is too cautious. VPBank is "
           "offering us eight billion dong over three years. If we do not scale, YouthWorks will "
           "eat our lunch in Binh Duong within a year.",
           "I have spoken informally to contacts at the provincial Department of Labour. They "
           "will back us if we commit to Long An. This is our moment. Let us not waste it.",
           "T.N."])
    email("Dr. Lan Phuong Nguyen", "T.N.; GBF Board distribution list", "Tue 6 Jan 2026, 08:02",
          "RE: We cannot afford to be timid",
          ["Thank you, T.N. I share the ambition. My concern is that our Dong Nai hub is not "
           "yet delivering, and scaling a model that is not working simply multiplies the "
           "problem. I have asked for an external review before we commit. I hope the Board will "
           "support that.",
           "Lan"])
    email("Anh Duc Pham (Development Director)", "GBF Board distribution list", "Tue 6 Jan 2026, 09:40",
          "RE: We cannot afford to be timid",
          ["For what it is worth, donors fund growth stories, not pauses. A 'stop and fix' "
           "narrative will cost us momentum with VPBank and others. We should commit publicly and "
           "solve the operational issues in parallel.",
           "Anh Duc"])
    _save(doc, "D08_Board_Email_Thread", "1_Client_Data_Room", "pdf")


def build_d10_competitor():
    doc = _doc("Competitor Brief: YouthWorks", "Prepared by the GBF Development team",
               classification="Client document D-10", ref="D-10", issued="November 2025")
    add_heading(doc, "Summary", 1)
    add_body(doc, "YouthWorks is a government-backed youth training initiative that launched in "
             "Binh Duong in November 2025. It offers a free six-week training course with an "
             "official government certificate. It has strong public funding and a recognised "
             "brand.")
    add_heading(doc, "How it compares to GBF", 1)
    add_table(doc, ["Feature", "YouthWorks", "GBF Skills Forward"], [
        ("Cost to young person", "Free", "Free"),
        ("Duration", "6 weeks", "16 weeks"),
        ("Paid internship", "No", "Yes (8 weeks)"),
        ("Government certificate", "Yes", "No"),
        ("Published placement rate", "Not disclosed", "71% (2024)"),
    ], widths=[2.3, 2.1, 2.1])
    add_heading(doc, "Assessment", 1)
    add_body(doc, "YouthWorks' key weakness is the absence of an internship, which is the part "
             "of our model most associated with real job outcomes. However, 'free and fast and "
             "certified' is attractive, and our field staff estimate that up to 40 per cent of "
             "our current waitlist would switch to YouthWorks if it had the capacity to take them.")
    add_body(doc, "The Development Director's view is that the threat is overstated because "
             "YouthWorks cannot place young people into jobs. Field staff in the border districts "
             "disagree. The truth is probably somewhere in between and worth testing against the "
             "data.", italic=True, size=10)
    _save(doc, "D10_Competitor_Brief_YouthWorks", "1_Client_Data_Room", "pdf")


def build_d11_staff_survey():
    doc = _doc("Staff Engagement Survey 2025", "Summary of results prepared by Human Resources",
               classification="Client document D-11", ref="D-11", issued="November 2025")
    add_heading(doc, "Method", 1)
    add_body(doc, "The annual anonymous survey was open for two weeks in October 2025. Fifty-two "
             "of sixty-seven full-time staff responded, a response rate of 78 per cent.")
    add_heading(doc, "Headline results", 1)
    add_table(doc, ["Statement", "Agree or strongly agree"], [
        ("My workload is sustainable", "39%"),
        ("I have the tools and data I need to do my job well", "44%"),
        ("I would recommend GBF as a place to work", "63%"),
        ("Our monitoring data is accurate", "31%"),
        ("Beneficiaries would benefit most from transport support", "72%"),
    ], widths=[4.2, 2.2])
    add_body(doc, "Field staff turnover during 2025 was 34 per cent, against a sector benchmark "
             "of approximately 22 per cent.")
    add_heading(doc, "Selected verbatim comments", 1)
    add_quote(doc, "I love the mission but I cannot keep working sixty-hour weeks. Something has to give.")
    add_quote(doc, "We are told to visit every employer monthly. With my caseload that is simply impossible.")
    add_quote(doc, "The data system is a mess. I do not trust our own placement numbers.")
    _save(doc, "D11_Staff_Engagement_Survey_2025", "1_Client_Data_Room", "pdf")


def build_d12_financials():
    doc = _doc("Financial Statements 2023-2025", "Audited summary prepared by Finance",
               classification="Client document D-12 (audited)", ref="D-12")
    add_heading(doc, "Statement of income by source", 1)
    add_body(doc, "All figures in VND billion.")
    add_table(doc, ["Source", "2023", "2024", "2025"], [
        ("Corporate foundations and CSR (incl. VPBank)", "6.9", "7.2", "7.6"),
        ("International NGOs and bilateral", "5.4", "5.1", "5.1"),
        ("Individual philanthropy", "1.9", "2.1", "2.2"),
        ("Government co-funding", "1.3", "1.5", "2.0"),
        ("Social enterprise surplus", "1.3", "1.7", "1.3"),
        ("Total income", "16.8", "17.6", "18.2"),
    ], widths=[3.4, 1.0, 1.0, 1.0])
    add_body(doc, "Of 2025 income, 78 per cent was restricted to specific programmes or "
             "provinces and 22 per cent was flexible core funding. The VPBank Foundation was the "
             "single largest funder at approximately 14 per cent of total income.")

    add_heading(doc, "Unit costs", 1)
    add_table(doc, ["Programme", "Cost per completer (VND)"], [
        ("Skills Forward (incl. staff time and internship support)", "12,800,000"),
        ("Pathway Digital", "6,200,000"),
    ], widths=[4.2, 2.2])

    add_heading(doc, "Note on the social enterprise", 1)
    add_body(doc, "GBF Staffing Solutions, established in 2019, places graduates into logistics "
             "and hospitality roles and earns a placement fee from employers. Its surplus "
             "cross-subsidises about 7 per cent of the budget. The auditors note the potential "
             "conflict of interest where the same organisation both advises young people and "
             "earns fees from placing them.")
    _save(doc, "D12_Financial_Statements_2023-2025", "1_Client_Data_Room", "pdf")


def build_interviews():
    doc = _doc("Interview Transcripts", "Field research conducted January 2026",
               classification="Client data room", ref="Transcripts ST/BN/DN/EP/VR/BD")
    add_body(doc, "The following transcripts have been lightly edited for clarity, not for "
             "brevity. Filler words and asides are retained where they show how people actually "
             "speak. Speaker names are shown in bold. The interviewer is a member of the "
             "engagement team. Cite conversations by reference code, for example (ST-01).",
             italic=True, size=10)

    def transcript(ref, name, role, exchanges):
        add_heading(doc, f"{ref}  -  {name}, {role}", 2)
        for speaker, line in exchanges:
            p = doc.add_paragraph()
            sr = p.add_run(f"{speaker}: ")
            sr.bold = True
            sr.font.size = Pt(10.5)
            br = p.add_run(line)
            br.font.size = Pt(10.5)
        doc.add_paragraph()

    transcript("ST-01", "Hung Vo", "Programmes Director", [
        ("Interviewer", "Thank you for making time. To start broadly - what is keeping you up "
         "at night right now?"),
        ("Hung", "Honestly? We are succeeding ourselves into a crisis. That sounds dramatic, but "
         "I mean it. There are 1,100 young people on the waitlist. Families call the hubs in "
         "tears. Mothers sit in the reception at Dong Nai and will not leave until someone "
         "gives them a date. And at the same time my field team is working sixty-hour weeks "
         "because the Board will not let us hire until there is a 'strategy'. So every week "
         "the queue grows and the people already inside the programmes get less attention."),
        ("Interviewer", "When you say the Board will not let you hire - is that a formal freeze?"),
        ("Hung", "It is a practical freeze. Anh Duc says we should wait for the VPBank decision "
         "before we add headcount, in case we need to restructure. I understand the logic on "
         "paper. In the field it means Minh and his colleagues are carrying caseloads that "
         "would get a social worker fired in any proper system. We are choosing growth talk "
         "over care, and I am tired of pretending that is strategy."),
        ("Interviewer", "Tell me more about Dong Nai specifically."),
        ("Hung", "Transport, first. The young people commute a long way to the industrial zones "
         "- twenty-five, sometimes forty kilometres - and they pay 150 to 200 thousand dong a "
         "month for buses that we do not cover. If the family is already stretched, that is "
         "the difference between staying and quitting. Internship drop-out hit 38 per cent "
         "last quarter against about 22 for the organisation. Second, employers. Some of our "
         "'partners' treat interns as free labour. Long hours, no learning plan, no contract "
         "conversation. I have names. I have reported them. Progress is slow because Development "
         "worries about the relationship."),
        ("Interviewer", "You mentioned a reporting disagreement with Anh Duc?"),
        ("Hung", "Yes. He wanted us to report 'completion' for Dong Nai using a softer "
         "denominator - basically anyone who finished classroom training, not internship "
         "completion. It would have made the numbers look less ugly. I refused. If we start "
         "that game with ourselves, we will eventually play it with VPBank, and then we "
         "deserve whatever happens."),
        ("Interviewer", "Should GBF expand to Long An?"),
        ("Hung", "Not until Dong Nai works. The Long An pilot lost about 60 per cent of its "
         "young people. Same pattern - distance, thin staffing, employers we barely know. "
         "There is a board member pushing Long An hard. I will not name him on the record, "
         "but his arguments sound more like political contacts than programme design. If we "
         "open another broken hub to please a Board faction, we will burn the VPBank money "
         "and the young people at the same time."),
        ("Interviewer", "And Pathway Digital - where do you stand?"),
        ("Hung", "It is cheaper per head, yes. On a spreadsheet it looks like the answer to "
         "scale. But the real completion rate is about 52 per cent, not the 68 we put in front "
         "of the donor. The digital team counted enrolments, not people who actually finished. "
         "I have asked them to correct it before it becomes a reputational problem. Also - and "
         "this matters - many of our young people do not have a laptop. Digital is not a "
         "neutral substitute for Skills Forward; it selects for those who already have more."),
        ("Interviewer", "If you could get one concrete thing from this consulting review, what "
         "would it be?"),
        ("Hung", "Permission - Board permission - to pause geographic expansion and fix quality "
         "for twelve months. Even if a couple of bad employers complain. Even if Anh Duc says "
         "we are 'losing momentum'. Momentum toward what? A larger waitlist and the same "
         "broken Dong Nai story, just with a bigger grant attached."),
    ])

    transcript("ST-02", "Thao Le", "Operations Director", [
        ("Interviewer", "We keep hearing different placement rates. Is the 71 per cent in the "
         "annual report accurate?"),
        ("Thao", "It depends what you mean by placed - and that is exactly the problem. If you "
         "include gig work at fifteen hours a week or more, then yes, on the public definition "
         "in D-01, you can get to around 71 for that reporting year. If you mean a formal "
         "labour contract with social insurance, it is closer to 63 per cent. VPBank knows we "
         "use the broader definition in the agreement language, but they have also asked, "
         "quietly, for formal to be shown separately. Anyone who quotes one number without "
         "saying which definition is either careless or selling something."),
        ("Interviewer", "What about the export Workflow B will see - should they expect 71?"),
        ("Thao", "No. The operational export is a different cut of time and a messier file. "
         "Tell them to define Completed versus Dropped, Formal versus Gig, and blank follow-ups "
         "before they celebrate or panic. I would rather a honest 65 with a footnote than a "
         "pretty 71 that collapses under audit."),
        ("Interviewer", "What does it actually cost to serve one young person?"),
        ("Thao", "Skills Forward is about 12.8 million dong per completer when you include staff "
         "time, not just the classroom budget. Pathway Digital is about 6.2 million, but its "
         "placement rate is lower and the completion story is softer, as Hung will tell you. "
         "People love the unit-cost comparison until you ask what quality you bought."),
        ("Interviewer", "Can GBF realistically reach 4,000 young people by 2028?"),
        ("Thao", "Only with the VPBank money and roughly twenty-five more staff - and that is "
         "optimistic. Recruitment is hard. We lose people to factory HR jobs and to YouthWorks, "
         "which pays less prestige but sometimes more certainty. Turnover was 34 per cent last "
         "year. You cannot scale a people-intensive model on a slide deck. You scale it with "
         "hiring pipelines, mentor benches, and supervisors who are not already drowning."),
        ("Interviewer", "The Executive Director has talked about an intake chatbot. Your view?"),
        ("Thao", "I think it is a distraction dressed as innovation. Our young people need a "
         "human voice on the phone - someone who can hear that the mother is crying, that the "
         "bus money is gone, that the accent is from the Delta. A chatbot that cannot parse "
         "that will just move the queue from WhatsApp into a CRM and call it efficiency. Fix "
         "Dong Nai transport before you buy software theatre."),
        ("Interviewer", "Anything else the team should pressure-test in the data?"),
        ("Thao", "Capacity. Look at hub burn rates and caseloads, not only placement. A scenario "
         "that clears a funding gap but breaks staff is not a scenario - it is a future "
         "resignation letter."),
    ])

    transcript("ST-03", "Linh Nguyen", "Monitoring & Evaluation Officer", [
        ("Interviewer", "How reliable is GBF's data, in your honest view?"),
        ("Linh", "Frankly, it is a mess - and I say that as the person who lives in it. We have "
         "duplicate beneficiary IDs with conflicting fields, dates in three different formats, "
         "and field staff who sometimes mark someone as placed after a single ride-hailing "
         "shift because the young person sounded hopeful on the phone. There are four of us "
         "covering about 800 young people in the current served population. We are not lazy; "
         "we are underbuilt. The public reports look clean because someone sanded them."),
        ("Interviewer", "If our analysts open the export, what should they check first?"),
        ("Linh", "Duplicates - de-duplicate carefully, and use last_updated if it is there. "
         "Normalise placed_90d before you count anything; blank is not the same as not placed. "
         "Look at Dong Nai costs for the third quarter - I think a month was entered twice. "
         "And there is at least one wage figure that is obviously a typo; it should be flagged "
         "in the notes. Also watch province labels: HCMC versus Ho Chi Minh City. People "
         "double-count that if they are not careful."),
        ("Interviewer", "What about young people with disabilities?"),
        ("Linh", "We report about six to seven per cent of beneficiaries with a disability "
         "flag. The population share among the young people we should be reaching is nearer "
         "nine. Their placement rate is about 45 to 50 per cent against roughly 70 on the "
         "public broad definition. Donors almost never see that breakdown, and when they do "
         "not ask, Development does not volunteer it. That is not a data problem only - it is "
         "a courage problem."),
        ("Interviewer", "Any last warning for the consulting team?"),
        ("Linh", "Do not trust a single rate. Ask for the definition, the denominator, the "
         "missingness, and who benefits if the number is high. If they cannot answer those "
         "four, they do not understand the number yet. And watch for pressure to make the "
         "picture look good earlier than the data can honestly support - that temptation never "
         "arrives wearing a villain costume."),
    ])

    transcript("ST-04", "Anh Duc Pham", "Development Director", [
        ("Interviewer", "How serious is the YouthWorks threat, from your side of the house?"),
        ("Anh Duc", "Overstated, in my view. They have speed and a free offer, but they have no "
         "real internship pathway. Our outcomes are better when young people stay the course. "
         "Yes, they are faster and we take longer - that is the trade. Young people who want a "
         "real job with insurance still come to us. The panic I hear from the field is real "
         "emotionally, but it is not a strategy case for abandoning our model."),
        ("Interviewer", "And the VPBank offer?"),
        ("Anh Duc", "It is the opportunity of a decade. Eight billion dong over three years if "
         "we hit 65 per cent and show a credible path to 4,000 young people by 2028. We should "
         "commit publicly. Hesitation kills donor confidence. Other organisations are already "
         "whispering that GBF is 'quality-obsessed and growth-shy'. That narrative sticks."),
        ("Interviewer", "Hung and others want to pause expansion and fix quality first. How do "
         "you hear that?"),
        ("Anh Duc", "I hear fear dressed as principle. Of course Dong Nai needs attention. Of "
         "course we should not invent numbers. But 'pause' without a growth path is how "
         "foundations quietly move money elsewhere. Bring me a plan that fixes quality and "
         "still scales. Do not bring me a sermon about integrity that ends with a smaller "
         "organisation and a disappointed Board."),
        ("Interviewer", "How should the Board think about measuring success under the grant?"),
        ("Anh Duc", "Practically. Donors and Boards need visibility before every story is "
         "fully baked. If we only talk outcomes when every file is perfect, we will always "
         "arrive late to the conversation. I want management dashboards that show momentum "
         "earlier - leading indicators, pipeline, early placement signals - so we can manage, "
         "not only audit. M&E sometimes treats that as heresy. I treat it as running an "
         "organisation. I am not asking anyone to invent results. I am asking us not to wait "
         "for perfect purity while competitors take the room."),
        ("Interviewer", "What would make you change your mind about aggressive scale?"),
        ("Anh Duc", "Hard evidence that we cannot hit 65 honestly under growth - not anecdotes, "
         "not one bad hub treated as destiny. If the data team shows that Aggressive growth "
         "forces selection bias or burns staff past recovery, I will listen. Until then, I "
         "will keep pushing for ambition. Someone has to."),
    ])

    transcript("ST-05", "Minh Tran", "Field Officer, Dong Nai", [
        ("Interviewer", "Describe your caseload as it actually feels day to day."),
        ("Minh", "There are four of us at the Dong Nai hub for roughly two hundred young people "
         "across the year - active cases sit closer to 120 at any one time, but the policy "
         "caseload target is about forty per officer and we are regularly over fifty. On paper "
         "we visit every employer monthly. That is simply not possible. Some weeks I choose "
         "between a crisis home visit and an employer check. Guess which one loses."),
        ("Interviewer", "You mentioned employer problems - can you be specific?"),
        ("Minh", "Two employers on our partner list exploit the interns - ten-hour days, no "
         "real learning duties, 'be grateful you are here' when someone asks for a contract. "
         "I reported them months ago with names and dates. Nothing meaningful happened. "
         "Development worries about losing the partner. I worry about losing the young person. "
         "Those worries are not equal."),
        ("Interviewer", "How does transport show up in your conversations?"),
        ("Minh", "Every week. A young person will say, quietly, that the bus is 180 thousand "
         "that month and their mother needs medicine. Then they stop coming. We call it "
         "drop-out. It is often arithmetic. We modelled quietly - just the field team - that "
         "a transport stipend of about 100 thousand dong a month would roughly halve our "
         "internship drop-out. Nobody has budgeted it because expansion stories are sexier."),
        ("Interviewer", "Is YouthWorks affecting your waitlist?"),
        ("Minh", "Last month they took about fifteen young people off our waitlist in the "
         "border districts. Families do not care about our theory of change when someone "
         "offers a free course next Tuesday. I do not blame them. I blame us for making them "
         "wait seven months while we argue about Board strategy in Ho Chi Minh City."),
        ("Interviewer", "What should the consultants not get wrong about Dong Nai?"),
        ("Minh", "Do not treat us as a 'underperforming hub' on a dashboard and then recommend "
         "more volume. We are underperforming because the model is under-resourced and some "
         "employers are bad actors. Fix those, then talk about scale. Also - please talk to "
         "the young people who left, not only the ones who finished and smiled for the donor "
         "photo."),
    ])

    transcript("BN-01", "Mai", "Programme graduate, Ho Chi Minh City", [
        ("Interviewer", "Mai, thank you for meeting us. Can you walk me through how you found "
         "GBF and what the waiting was like?"),
        ("Mai", "I heard about Skills Forward from a neighbour whose son did it the year "
         "before. I applied and then I waited. Five months. At month three I almost gave up - "
         "I was helping at a small coffee place and the owner said he might train me himself. "
         "My mother said wait for GBF because of the insurance. So I waited. When the call "
         "finally came I cried on the bus home. It sounds silly, but the waiting is part of "
         "the story. People talk about the training. They forget the months of nothing."),
        ("Interviewer", "What was the training like once you started?"),
        ("Mai", "Good, mostly. I learned Excel properly, customer service, how to write a CV "
         "that does not sound like a child wrote it. The teachers were patient. The group "
         "helped each other. I liked that. It felt like someone was finally taking us "
         "seriously as workers, not as charity cases."),
        ("Interviewer", "And the internship?"),
        ("Mai", "That was harder. It was far from home - about 170 thousand dong a month in "
         "bus fares that I paid myself. Some days I skipped lunch to make the money stretch. "
         "But the company was decent. They explained the work. They did not shout. After the "
         "internship they offered me a real contract, 5.8 million dong a month, with social "
         "insurance. I was lucky. I know I was lucky."),
        ("Interviewer", "Lucky compared to whom?"),
        ("Mai", "My friend in the Dong Nai group. She quit because the employer made them work "
         "ten-hour days for nothing useful - just sorting boxes and being told to smile. She "
         "said GBF never came to check. When I told my caseworker, they sighed like they had "
         "heard it before. That sigh stayed with me."),
        ("Interviewer", "Would Pathway Digital have suited you instead?"),
        ("Mai", "No. I do not have a laptop and my phone is old. The screen cracks and Zoom "
         "freezes. Skills Forward was right for me because I needed a place to go and people "
         "to practise with. Online sounds modern. For me it would have been lonely and "
         "impossible."),
        ("Interviewer", "If GBF asked what to fix first, what would you say?"),
        ("Mai", "Help with transport during internship. And check the employers. A certificate "
         "means nothing if the workplace teaches you that you do not matter."),
    ])

    transcript("BN-02", "Tuan", "Withdrew from the programme, Dong Nai", [
        ("Interviewer", "Tuan, we know this might be uncomfortable. Can you tell us why you "
         "left the programme?"),
        ("Tuan", "The training was fine - I am not angry about the classroom. The internship "
         "was miserable. No contract conversation, ten-hour days, and the boss told us to be "
         "grateful we were even there. Grateful for what? Free work? My mother is ill. I help "
         "at her coffee stall when I can. I could not keep paying for buses and losing hours "
         "at the stall while a man shouted that I should smile more. So I left. GBF called it "
         "drop-out. I call it choosing my mother over a bad placement."),
        ("Interviewer", "Did anyone from GBF visit the workplace before or during?"),
        ("Tuan", "At the start, maybe once, briefly. After that, no. When I complained to my "
         "caseworker he said he would 'note it'. Nothing changed. People in meetings talk "
         "about numbers. Nobody asked me what the bus cost or what the boss said when I asked "
         "about insurance. I still want a factory job with proper social insurance. I did not "
         "leave because I am lazy. I left because the internship was not a pathway; it was a "
         "discount worker scheme with a GBF logo on it."),
        ("Interviewer", "What would have kept you in?"),
        ("Tuan", "Employers who respect us. A real check from GBF that is not just a photo for "
         "Facebook. And some help with transport - even 100 thousand a month would have "
         "changed the calculation. Also honesty. If the placement is bad, say it is bad. Do "
         "not tell us we are lucky while we are being used."),
        ("Interviewer", "Have you looked at YouthWorks since?"),
        ("Tuan", "I went to one information session. It was short and free. No internship. I "
         "did not join yet. I am waiting to see if GBF cleans up. My mother still asks when I "
         "will have insurance. I do not know what to tell her."),
        ("Interviewer", "Anything you want the Board to hear, in your words?"),
        ("Tuan", "Stop counting me as a failure metric and start counting the bus fare and the "
         "bad bosses. If you expand before you fix that, you will just create more people like "
         "me - quiet, gone, and still needing work."),
    ])

    transcript("BN-03", "Lan", "Graduate with a disability, Ho Chi Minh City", [
        ("Interviewer", "Lan, was the programme accessible for you in practice?"),
        ("Lan", "At first, no. The building had no ramp. I had to be carried up two steps by "
         "classmates, which is humiliating even when people are kind. After I complained - and "
         "I had to complain more than once - they moved the class downstairs. That should not "
         "have required a campaign. My mentor was kind, genuinely, but she did not know much "
         "about disability employment rights. She kept saying 'stay positive'. Positivity does "
         "not install a ramp or change an employer's mind."),
        ("Interviewer", "How has the job search gone since you finished?"),
        ("Lan", "Still looking after four months. Interviews where they smile and say they "
         "will call, and then silence. One HR person asked, in front of others, whether I "
         "would 'slow the team down'. GBF counts me as 'in progress' on their list. I "
         "understand why they need categories. Still, it feels like I am a statistic with a "
         "polite label, not a person they are fighting for."),
        ("Interviewer", "Did the programme talk about disability and work openly?"),
        ("Lan", "Very little. Soft skills, CV, interview practice - useful, but generic. Nobody "
         "role-played how to answer a discriminatory question. Nobody connected me to "
         "employers who already hire people with disabilities. Mentoring helped emotionally. "
         "It did not open doors."),
        ("Interviewer", "What would good support look like?"),
        ("Lan", "Access sorted before day one, not after a complaint. Mentors trained on "
         "rights, not only encouragement. Employers screened for accessibility, the same way "
         "you should screen for exploitation. And stop hiding our placement rate inside the "
         "overall number. If it is lower - and Linh says it is - say so and fix it."),
        ("Interviewer", "Would you recommend GBF to another young person with a disability?"),
        ("Lan", "I would recommend it with warnings. Better than nothing. Not yet good enough "
         "to be proud of."),
    ])

    transcript("BN-04", "Hang", "On the waitlist, Binh Duong", [
        ("Interviewer", "Hang, how long have you been waiting, and what has that been like?"),
        ("Hang", "Seven months. I check my phone every morning like a superstition. Nothing. "
         "My father asks weekly whether I am 'still waiting for that charity course'. He wants "
         "me to take any factory overtime. I tell him GBF means a real pathway. Some weeks I "
         "believe myself. Some weeks I do not."),
        ("Interviewer", "Have other programmes approached you?"),
        ("Hang", "YouthWorks offered a six-week course, free, starting quickly. No internship "
         "at the end. I went to the briefing. It felt rushed - smile, certificate, goodbye. I "
         "stayed on the GBF list because I want a job with social insurance, not only a short "
         "class. But every month I wait, YouthWorks looks more reasonable. Waiting is also a "
         "decision cost."),
        ("Interviewer", "What would make you finally switch?"),
        ("Hang", "If GBF cannot give me a start date within the next two months, I will take "
         "YouthWorks and keep looking for insurance later. I cannot be loyal to a queue "
         "forever. Loyalty does not pay rice."),
        ("Interviewer", "What do you actually need from a programme?"),
        ("Hang", "Something that ends in a contract, or at least a serious employer interview. "
         "Transport help if the workplace is far. And honesty about timing. Do not tell me "
         "'soon' for seven months. Say the number on the waitlist and where I sit. Treat me "
         "like an adult."),
        ("Interviewer", "Any message for the people designing GBF's next three years?"),
        ("Hang", "If you scale without shortening the wait, you are scaling disappointment. I "
         "am not against growth. I am against growth that leaves people like me refreshing a "
         "blank inbox."),
    ])

    transcript("DN-01", "Ms. Huong", "Programme Officer, VPBank Foundation", [
        ("Interviewer", "What does VPBank Foundation need to see to renew and expand support?"),
        ("Huong", "Three things, stated simply. First, a placement rate of at least 65 per "
         "cent at ninety days under a definition we can audit. Second, a credible path to "
         "4,000 young people by 2028 - not a slogan, a resourced plan. Third, a monitoring "
         "framework we can trust, including the ugly numbers. We are a foundation attached to "
         "a bank. We understand risk. We do not understand surprise."),
        ("Interviewer", "Does gig work count toward the 65 per cent?"),
        ("Huong", "If it is stable and at least fifteen hours a week, the agreement language "
         "allows it. But we would strongly prefer formal employment reported separately. "
         "Dual reporting protects everyone. A single blended rate that hides fragility will "
         "eventually damage the relationship."),
        ("Interviewer", "What if GBF asked to slow geographic expansion and fix quality first?"),
        ("Huong", "We would listen. We could consider a smaller, quality-focused grant structure "
         "for a period. But the Board would need to bring us a credible theory of change - "
         "what will be fixed, by when, with which leading indicators - not simply 'give us "
         "time'. Other organisations are competing for this money. Patience without a plan is "
         "not a proposal."),
        ("Interviewer", "What would worry you most in a board recommendation from this team?"),
        ("Huong", "Aggressive scale with no answer for Dong Nai. Or a recommendation that "
         "treats 65 per cent as a branding target rather than an operational truth. Bring us "
         "ambition and honesty in the same document. We can work with that. And please read "
         "the agreement carefully before you invent a cleverer measurement story than the one "
         "we already negotiated."),
    ])

    transcript("EP-01", "Mr. Khang", "HR Manager, gold-tier employer partner", [
        ("Interviewer", "How do GBF interns perform in your operation?"),
        ("Khang", "Mixed, and I will not pretend otherwise. The best are excellent - we hired "
         "twelve last year into longer contracts. They show up, they learn the floor, they ask "
         "questions. The weakest lack basic punctuality and seem surprised that logistics is "
         "physical work. GBF prepares them well on soft skills and confidence. Factory "
         "discipline and stamina are thinner. That gap matters on a 5 a.m. shift."),
        ("Interviewer", "Would you take more interns if GBF asked?"),
        ("Khang", "If they were better pre-screened, yes. We get too many who want an office "
         "job and discover on day two that the role is warehouse and routing. That wastes "
         "everyone's time. Also - and I will be frank - it is awkward that GBF Staffing "
         "Solutions then charges us a fee to hire the same young people permanently. It feels "
         "like we are paying twice: once in training goodwill, once in a placement fee. I "
         "still partner with GBF because the best graduates are worth it. But do not ask me to "
         "call that structure elegant."),
        ("Interviewer", "Have you seen differences by hub - for example Dong Nai versus HCMC?"),
        ("Khang", "I do not track hubs the way GBF does. I track who arrives ready. When "
         "someone arrives exhausted from a long commute and already bitter about money, the "
         "internship fails before skill is tested. That is not only a training issue."),
        ("Interviewer", "What should GBF stop doing, from an employer's seat?"),
        ("Khang", "Stop sending volume to look busy. Send fit. And visit the floor after week "
         "two, not only at signing. Partners notice when oversight is ceremonial."),
    ])

    transcript("EP-02", "Ms. Trang", "Operations manager, logistics subcontractor (anonymous)", [
        ("Interviewer", "Thank you for speaking with us off the record. How does your firm "
         "actually use GBF interns?"),
        ("Trang", "Look, I will be straight with you, because the polite version helps nobody. "
         "We are short-staffed in peak season. The interns help with sorting and packing. We "
         "pay the legal minimum stipend, but the work is repetitive and the hours are long. "
         "There is no beautiful learning journey. GBF visited once at the start - handshake, "
         "photo, gone. We have not seen a caseworker since. I am not proud of how we use them. "
         "If GBF stopped sending people, we would manage with overtime and contractors. The "
         "interns are convenient, not sacred."),
        ("Interviewer", "Do the young people complain?"),
        ("Trang", "Some do, quietly. Some quit. We replace them. That sentence should bother "
         "you. It bothers me when I say it out loud. Inside the business, it is normalised."),
        ("Interviewer", "Would better monitoring from GBF change your behaviour?"),
        ("Trang", "Yes. If someone actually checked hours, duties, and whether we were teaching "
         "anything, we would clean up. Not because we became saints overnight - because we "
         "would not risk the partner channel. Right now there is no risk. No risk means no "
         "change."),
        ("Interviewer", "Would you take more interns under current arrangements?"),
        ("Trang", "Only if someone checked that we were treating them properly. Otherwise you "
         "are asking me to participate in a fiction - that this is always 'opportunity'. "
         "Sometimes it is just cheap labour with a programme name. You should decide whether "
         "that is acceptable before you scale."),
        ("Interviewer", "Anything else the Board should hear from your side of the market?"),
        ("Trang", "Gold-tier partners will give you nice quotes. Ask the subcontractors too. "
         "That is where the harm hides."),
    ])

    transcript("VR-01", "Chi Hoa", "Volunteer mentor, three years", [
        ("Interviewer", "Is the mentoring model sustainable from where you sit?"),
        ("Hoa", "Barely. Volunteers burn out. I mentor four young people, which is too many if "
         "you are doing it properly - weekly calls, emergency messages, sitting with someone "
         "while they cry about a boss or a parent. We get a two-hour orientation and then we "
         "are on our own. When a case is light, it is rewarding. When it is heavy, I feel "
         "unqualified and still responsible. That combination empties people."),
        ("Interviewer", "What kind of 'heavy' are we talking about?"),
        ("Hoa", "Domestic violence. Debt collectors. A young person who messaged me at midnight "
         "about not wanting to wake up. I am a mentor, not a clinician. I called the field "
         "officer and stayed on the phone until morning. We need professionals in the chain. "
         "Volunteers can walk beside someone. We cannot be the entire safety system."),
        ("Interviewer", "Could volunteers replace paid field staff to save cost while scaling?"),
        ("Hoa", "No. Absolutely not. Anyone who suggests that has not taken those midnight "
         "calls. Volunteers support the work; we cannot be the work. If the Board tries to "
         "replace caseworkers with mentors to hit a headcount target, you will get a quieter "
         "spreadsheet and louder harm."),
        ("Interviewer", "What would make mentoring healthier?"),
        ("Hoa", "Fewer young people per mentor. Better supervision. Clear escalation paths. "
         "And please stop listing inactive mentors as if their hours still count - I have seen "
         "names on lists of people who left months ago. It makes the capacity look stronger "
         "than it is."),
        ("Interviewer", "Would you still encourage a friend to volunteer?"),
        ("Hoa", "Yes, with eyes open. It matters. It is also unpaid emotional labour next to a "
         "professional gap. Be honest in the recruitment poster."),
    ])

    transcript("BD-01", "Mr. T.N.", "Member of the GBF Board", [
        ("Interviewer", "Where do you stand on the strategic choice in front of the Board?"),
        ("T.N.", "We must accept the VPBank money and scale - including Long An. I have "
         "political and provincial support lined up. If we hesitate, we lose momentum and we "
         "lose donors. The region is watching whether GBF is serious about growth or only "
         "serious about workshops in District 1. The 'pause and fix' talk worries me. It "
         "sounds virtuous. It often becomes delay."),
        ("Interviewer", "What do you say to staff who argue Dong Nai is not ready?"),
        ("T.N.", "Fix Dong Nai in parallel. Do not hold the whole organisation hostage to one "
         "hub's problems. Every expanding NGO has a weak site. That is not a reason to refuse "
         "a once-in-a-decade grant. Hire, train, tighten employer standards - yes. Hide from "
         "growth - no."),
        ("Interviewer", "Some voices say Long An expansion is influenced by personal networks."),
        ("T.N.", "People always say that when they dislike a geography. Long An is part of the "
         "economic zone story. Young people there need pathways too. If the pilot had issues, "
         "improve the model and go again. Retreating to HCMC comfort is not equity."),
        ("Interviewer", "What would you need from this consulting team to support their "
         "recommendation?"),
        ("T.N.", "A plan that says yes to VPBank with eyes open - staffing numbers, timeline to "
         "4,000, and how 65 per cent stays honest. If they come back with 'slow down and hope "
         "the foundation waits', they will find a chilly Board room. Ambition is not the "
         "enemy of integrity. Fear is."),
        ("Interviewer", "Any red line for you personally?"),
        ("T.N.", "Publicly refusing the grant without a better funded alternative. That would "
         "be organisational self-harm dressed as principle."),
    ])

    _save(doc, "20_Interview_Transcripts", "1_Client_Data_Room", "pdf")

# ===========================================================================
# 2. WORKSTREAM A
# ===========================================================================

def build_wsa_brief():
    doc = _doc("Workflow A - Impact Strategy", "Engagement brief and requirements",
               classification="Participant material - Workflow A", ref="MOM-2026-WSA-01")
    add_heading(doc, "Purpose of this workflow", 1)
    add_body(doc, "Workflow A owns direction. You make sense of what stakeholders say, frame "
             "the problem, decide which questions matter, develop strategic options, and build "
             "the story the Board will hear. Workflow B holds the operational record. You cannot "
             "look the numbers up yourself; you must ask.")

    add_heading(doc, "Hard constraint on data", 1)
    add_body(doc, "You do not have access to GBF_Datasets.xlsx or the data dictionary. Do not "
             "ask facilitators for those files. You will receive a Data Estate Brief from B "
             "(what exists, in plain language) - not the raw file. When you need a quantitative "
             "answer, file a Demand Brief, then formal Analysis Requests. Working out what to "
             "ask is part of the exercise. Vague demands get vague plans.", bold=True)

    add_heading(doc, "How the exchange with B works", 1)
    add_body(doc, "B is not idle while you frame. In Weeks 1-3 they profile and clean the export, "
             "write a Data Estate Brief for you, and build a working dashboard. Your job in Week 3 "
             "is to read that brief, file a Demand Brief (ranked decision needs from your A1 "
             "hypotheses), then file precise R-xx requests. B responds with an Analysis Plan "
             "(they own the method) and Findings Memos. Do not delay Demand / R-01 until options "
             "are finished - B needs early demand to plan.")

    add_heading(doc, "Your responsibilities", 1)
    add_bullet(doc, "Synthesise interview transcripts into clear, cited insights.")
    add_bullet(doc, "Frame the problem with a MECE issue tree and testable hypotheses.")
    add_bullet(doc, "Read B's Data Estate Brief; file a Demand Brief that states what you need "
               "for the Board decision (not column names).")
    add_bullet(doc, "File Analysis Requests that turn Demand items into concrete R-xx asks.")
    add_bullet(doc, "Ground recommendations in specific voices from the transcripts.")
    add_bullet(doc, "Develop two or three strategic options and a single recommendation.")
    add_bullet(doc, "Sit with the costs of your preferred choice (trade-off reflection).")
    add_bullet(doc, "Own the executive summary and the narrative of the board presentation.")

    add_heading(doc, "Deliverables, owners and deadlines", 1)
    add_table(doc, ["#", "Deliverable", "Template", "Due"], [
        ("A1", "Problem statement and issue tree", "Templates/Problem_Statement_and_Issue_Tree.docx", "Wed 5 Aug"),
        ("A1b", "Stakeholder map", "Templates/Stakeholder_Map.docx", "Thu 6 Aug"),
        ("C", "Information credibility matrix (shared)", "4_Shared_Toolkit/Templates/Information_Credibility_Matrix.docx", "Thu 6 Aug"),
        ("DMD", "Demand Brief (to Workflow B)", "4_Shared_Toolkit/Templates/Demand_Brief.docx", "Tue Week 3"),
        ("A2", "Interview synthesis", "Templates/Interview_Synthesis.docx", "Wed 12 Aug"),
        ("A3", "Empathy and journey maps", "Templates/Empathy_and_Journey_Map.docx", "Wed 12 Aug"),
        ("R", "Analysis Requests (minimum three)", "4_Shared_Toolkit/Templates/Analysis_Request.docx", "Weeks 3-5"),
        ("A4", "Strategic options", "Templates/Strategic_Options.docx", "Thu 20 Aug"),
        ("A5", "Trade-off reflection", "Templates/Tradeoff_Reflection.docx", "Wed 26 Aug"),
        ("A6", "Validation protocol (with WF-B)", "Templates/Validation_Protocol.docx", "Wed 26 Aug"),
        ("A7", "Recommendation revision log", "Templates/Recommendation_Revision_Log.docx", "Thu 27 Aug"),
    ], widths=[0.5, 2.4, 2.8, 0.9])

    add_heading(doc, "Detailed requirements", 1)
    add_body(doc, "A1 - Problem statement and issue tree.", bold=True)
    add_bullet(doc, "Write a precise problem / decision statement in your own words. Do not copy "
               "the engagement letter. The Board hired you partly to frame the decision.")
    add_bullet(doc, "Build an issue tree that is mutually exclusive and collectively exhaustive. "
               "Invent the branches from the evidence - the template does not prescribe them.")
    add_bullet(doc, "Write at least four hypotheses. Each hypothesis should imply a question you "
               "could ask Workflow B.")

    add_body(doc, "Demand Brief and Analysis Requests.", bold=True)
    add_bullet(doc, "Wait for B's Data Estate Brief before locking Demand - use it to aim asks, "
               "not to invent field names.")
    add_bullet(doc, "File the Demand Brief in Week 3 with ranked decision needs linked to A1 "
               "hypotheses. Demand is 'what would change the recommendation', not 'send columns'.")
    add_bullet(doc, "Before final options, complete at least three requests with Findings Memos returned.")
    add_bullet(doc, "File R-01 once Demand and B's Analysis Plan for that item exist - do not wait "
               "for a polished deck.")
    add_bullet(doc, "Each request states: the decision it informs, the hypothesis, what a useful "
               "answer would look like, and urgency. Do not prescribe field names you have not been given.")
    add_bullet(doc, "Ask what would change your mind - not everything that looks interesting. "
               "Use the official templates; chat messages are not the record.")
    add_bullet(doc, "Cite Request IDs (R-01, R-02...) on every quantitative claim in A4 and later.")
    add_bullet(doc, "Save DEB, Demand, Plans, Requests and Memos as PDF in A_B_Exchange/ "
               "(see Programme Guideline). Markers assess the exchange from that folder.")

    add_body(doc, "A2 - Interview synthesis.", bold=True)
    add_bullet(doc, "Summarise by theme, not by person. Flag contradictions between sources.")
    add_bullet(doc, "Every insight cites at least one transcript reference.")

    add_body(doc, "A3 - Empathy and journey maps.", bold=True)
    add_bullet(doc, "Choose one beneficiary archetype grounded in BN-01 to BN-04.")
    add_bullet(doc, "Map the journey from awareness to a job at ninety days. Attach a quote or "
               "a Findings Memo cite to each major pain point.")

    add_body(doc, "A4 - Strategic options.", bold=True)
    add_bullet(doc, "Present two or three genuinely different options.")
    add_bullet(doc, "For each: rationale, cost implication (from B), main risks, who wins and loses.")
    add_bullet(doc, "Recommend one. State what would change your mind.")

    add_body(doc, "A5 - Trade-off reflection.", bold=True)
    add_bullet(doc, "Answer the open prompts in the template honestly. There is no checklist of "
               "named dilemmas to tick off.")
    add_bullet(doc, "Make clear what your recommendation costs - and for whom.")

    add_heading(doc, "How Workflow B will challenge you", 1)
    add_body(doc, "Before you submit strategic options or final board materials, Workflow B will "
             "check that every number cites a Request or Findings ID, that options are financially "
             "realistic, and that you have not ignored uncomfortable findings. Those gated "
             "submissions also require their online confirmation on the handoff checklist "
             "(Guideline Section 8).")
    _save(doc, "WSA_Engagement_Brief", "2_Workflow_A_Impact_Strategy", "pdf")


def build_wsa_playbook():
    doc = _doc("Design-Thinking Playbook", "Facilitated activities for Workflow A",
               classification="Participant material - Workflow A", ref="MOM-2026-WSA-02")
    add_body(doc, "This playbook gives you a run sheet for each design-thinking activity. Every "
             "activity must connect to evidence. Design thinking interprets the evidence; it does "
             "not replace it. When you need a number, write an Analysis Request rather than guessing.")

    for title, when, dur, steps, output in [
        ("Stakeholder mapping", "Week 2", "90 minutes",
         ["Brainstorm every party affected by the decision, one per sticky note.",
          "Cluster them, then plot on a power/interest grid.",
          "Star the five you most need to understand and record why."],
         "A stakeholder map and a shortlist for deeper study."),
        ("Problem framing", "Week 2", "120 minutes",
         ["Each member writes 'the decision the Board must make is...' independently.",
          "Share, challenge, and agree a single problem / decision statement in your own words "
          "(not a paraphrase of the engagement letter).",
          "Build an issue tree from the evidence; generate hypotheses that imply asks for Workflow B."],
         "A problem statement, issue tree and hypotheses (deliverable A1)."),
        ("First demand and request cycle", "Week 3", "90 minutes",
         ["Read Workflow B's Data Estate Brief (catalog only - if they offered the raw file, refuse it).",
          "From A1 hypotheses, draft the Demand Brief: ranked decision needs and what would change your mind.",
          "Save Demand Brief PDF to A_B_Exchange/; wait for B's Analysis Plan on the top items.",
          "Complete Analysis Request R-01 for the top Demand item; log it; hand to B."],
         "DMD + R-01 in A_B_Exchange/; waiting for Analysis Plan / FM-01."),
        ("Empathy mapping", "Week 3", "90 minutes",
         ["Choose one beneficiary archetype grounded in BN-01 to BN-04.",
          "Fill the four quadrants - says, thinks, does, feels - using cited quotes only.",
          "List what you still do not know, and which of those gaps need a request to B."],
         "An evidence-based empathy map (part of deliverable A3)."),
        ("Journey mapping", "Week 3", "120 minutes",
         ["Map the stages from awareness to ninety-day employment.",
          "Mark the emotional highs and lows and the pain points.",
          "Attach a quote or a Findings Memo cite to each pain point."],
         "A current-state journey map with prioritised pain points."),
        ("Ideation - How Might We and Crazy Eights", "Week 4", "120 minutes",
         ["Turn each priority pain point into a 'How might we...' question.",
          "Generate ideas quickly and without judgement, then sketch eight variations of the best.",
          "Dot-vote; check surviving ideas against what the transcripts actually say."],
         "A shortlist of concepts to develop into options."),
        ("Prioritisation with Findings", "Week 4", "60 minutes",
         ["Plot options on an impact-versus-feasibility grid.",
          "Ask Workflow B to challenge placements using Findings Memos, not raw tables.",
          "Record the assumptions behind every placement."],
         "A prioritisation matrix feeding deliverable A4."),
    ]:
        add_heading(doc, title, 2)
        add_meta_table(doc, [("When", when), ("Duration", dur)])
        add_body(doc, "Run sheet:", bold=True, space_after=2)
        for s in steps:
            add_number(doc, s)
        add_body(doc, f"Output: {output}", italic=True, size=10)
    _save(doc, "WSA_Design_Thinking_Playbook", "2_Workflow_A_Impact_Strategy", "pdf")


# ===========================================================================
# 3. WORKSTREAM B
# ===========================================================================

def build_wsb_brief():
    doc = _doc("Workflow B - Operations & Analytics", "Engagement brief and requirements",
               classification="Participant material - Workflow B", ref="MOM-2026-WSB-01")
    add_heading(doc, "Purpose of this workflow", 1)
    add_body(doc, "Workflow B owns the operational evidence. You hold the datasets and the data "
             "dictionary. You clean, analyse, and translate findings into language Workflow A can "
             "use. Success is not a dump of pivots; it is a clear answer to a question A actually asked "
             "- delivered fast because you prepared the data before the ask arrived.")

    add_heading(doc, "Hard constraint on sharing", 1)
    add_body(doc, "Do not share GBF_Datasets.xlsx or the data dictionary with Workflow A. "
             "You may share a Data Estate Brief (catalog in plain language), an Analysis Plan "
             "(how you will answer), Findings Memos and curated charts. If A asks for 'the file', "
             "refuse and ask what decision the number needs to support.", bold=True)

    add_heading(doc, "Three modes of work (do not wait idle)", 1)
    add_body(doc, "You do not need A's permission to understand the data. You do need Demand / "
             "Requests before you push decision numbers into A's narrative.")
    add_table(doc, ["Mode", "When", "What you do"], [
        ("Foundation (proactive)", "Weeks 1-3",
         "Read the dictionary; profile and clean (B1 / DB-1); write the Data Estate Brief for A; "
         "build a working dashboard (B2 / DB-2)."),
        ("Planning (proactive + joint)", "Week 3 onward",
         "Read A's Demand Brief; write Analysis Plans you own; then answer formal R-xx with Findings Memos."),
        ("Challenge (joint)", "Weeks 4-6",
         "Evidence-check A's options (B4); run funding scenarios (B5) once options exist."),
    ], widths=[1.6, 1.2, 3.6])
    add_body(doc, "If no Demand Brief has arrived by mid-Week 3, nudge A. Do not invent a "
             "strategy for them - ask what decision is stuck. Do not dump the dictionary on A "
             "instead of a proper Data Estate Brief.")

    add_heading(doc, "Your responsibilities", 1)
    add_bullet(doc, "Own and apply the data dictionary; profile and clean the export.")
    add_bullet(doc, "Brief A with a Data Estate Brief (catalog / limits - never the raw workbook).")
    add_bullet(doc, "After A's Demand Brief, write Analysis Plans that show how you will answer.")
    add_bullet(doc, "Maintain the Request Log; answer Analysis Requests on time with Findings Memos.")
    add_bullet(doc, "Translate findings into plain language with caveats and definitions.")
    add_bullet(doc, "Build a working dashboard (B only) and pull curated charts into Findings Memos.")
    add_bullet(doc, "Model funding scenarios behind any recommendation A is considering, "
               "including staff, mentor and finance capacity (hub_capacity_jan2026).")
    add_bullet(doc, "Give an honest assessment of placement under clear definitions.")
    add_bullet(doc, "Maintain the assumption log, risk register and version control.")

    add_heading(doc, "Deliverables, owners and deadlines", 1)
    add_table(doc, ["#", "Deliverable", "Format", "Due"], [
        ("DEB", "Data Estate Brief (to Workflow A)", "4_Shared_Toolkit/Templates/Data_Estate_Brief.docx → PDF", "Mon Week 3"),
        ("B1", "Data quality report (internal to B)", "Templates/Data_Quality_Report.docx -> PDF", "Tue 11 Aug"),
        ("AP", "Analysis Plans (linked to Demand / R-xx)", "4_Shared_Toolkit/Templates/Analysis_Plan.docx", "From Week 3"),
        ("FM", "Findings Memos (responses to R-01+)", "4_Shared_Toolkit/Templates/Findings_Memo.docx", "From Week 3"),
        ("B2", "Performance dashboard (working file)", "GBF_Performance_Dashboard.xlsx", "Fri 14 Aug"),
        ("B3", "Analysis summary (tied to Requests answered)", "Templates/Analysis_Summary.docx -> PDF", "Wed 19 Aug"),
        ("B4", "Evidence-check memo on A's options", "Templates/Evidence_Check_Memo.docx -> PDF", "Fri 21 Aug"),
        ("B5", "Funding scenarios", "Toolkit workbook, 'Funding Scenarios' tab", "Thu 27 Aug"),
    ], widths=[0.5, 2.5, 2.7, 0.9])

    add_heading(doc, "Analysis backlog", 1)
    add_body(doc, "DB-1 and DB-2 are mandatory foundation. The Data Estate Brief is how A learns "
             "what you hold without seeing the file. Everything else follows Demand: when A states "
             "a decision need, you write the Analysis Plan, then analyse, then return a Findings "
             "Memo. If you are unsure how to approach a demand, ask your Buddy for coaching "
             "questions - not for the answer.")

    add_body(doc, "DB-1 - Data quality audit (foundation).", bold=True)
    add_bullet(doc, "Profile every table for missing values, duplicates and format problems.")
    add_bullet(doc, "Use WSB_Data_Dictionary.pdf as your field guide. Investigate anomalies; "
               "do not delete silently.")
    add_bullet(doc, "Record every cleaning decision in the assumption log. Share with A only the "
               "material issues that affect a decision (via Findings Memo), not the full cleaning log.")

    add_body(doc, "DB-2 - Performance dashboard (foundation).", bold=True)
    add_bullet(doc, "Use pivot tables to show placement by hub, programme and gender.")
    add_bullet(doc, "Flag any rate that looks material under the definition you state.")
    add_bullet(doc, "Keep this as a B working file. Export only curated views into Findings Memos "
               "when A asks.")

    add_body(doc, "On-request analysis (no fixed recipe).", bold=True)
    add_bullet(doc, "After A's Demand Brief, write an Analysis Plan you own, then answer each "
               "formal R-xx with a Findings Memo in A_B_Exchange/.")
    add_bullet(doc, "Your job: restate their decision question, choose tables from the dictionary, "
               "compute carefully, state definitions and caveats.")
    add_bullet(doc, "Write Analysis Plans against Demand (before R-xx). Do not invent Plans with "
               "no Demand, and do not share raw tables. If stuck on method, ask Buddy - they will "
               "nudge, not solve.")
    add_bullet(doc, "Chat acknowledgements are fine; the assessable answer is the FM-xx PDF.")

    add_body(doc, "Funding scenarios (after draft options).", bold=True)
    add_bullet(doc, "Run once A has sketched options or asks what growth costs (deliverable B5). "
               "Stress-test each scenario against mission fit, Board ambition, and internal "
               "capacity (staff FTE, mentors, monthly burn) - not cost alone.")
    add_bullet(doc, "Use audited costs (D-12) and funder terms (D-03); state every assumption.")
    add_bullet(doc, "Judge whether each scenario can meet placement conditions under an honest "
               "definition you state.")

    add_heading(doc, "How Workflow A will challenge you", 1)
    add_body(doc, "Before gated submissions (options, final report, board pack), Workflow A will "
             "check that you answered their questions in plain language, that you did not share "
             "the raw file, that definitions and caveats are stated, and that unanswered requests "
             "are logged. Those submissions also require their online confirmation on the handoff "
             "checklist (Guideline Section 8).")
    _save(doc, "WSB_Engagement_Brief", "3_Workflow_B_Operations_and_Analytics", "pdf")


def build_wsb_data_dictionary():
    doc = _doc("Data Dictionary", "Operational export held by Workflow B",
               classification="Workflow B only - do not share with Workflow A",
               ref="MOM-2026-WSB-DD")
    add_body(doc, "This dictionary describes the nine tables in GBF_Datasets.xlsx. It is for "
             "Workflow B only. Do not forward this file or the workbook to Workflow A.")

    add_heading(doc, "Tables", 1)
    add_table(doc, ["Table", "Approx. rows", "Purpose"], [
        ("beneficiary_registry", "~802 rows / 800 unique IDs", "Master list (2 intentional duplicate IDs); matches guideline served 2025"),
        ("employment_outcomes", "~700", "Employment status ~90 days after completion/drop"),
        ("programme_attendance", "~5,000", "Session-level attendance (incl. BM mentor circles)"),
        ("survey_outcomes_2025", "~500", "End-of-programme satisfaction survey"),
        ("volunteer_hours", "~1,700", "Mentor hours logged (check status)"),
        ("hub_costs_2025", "60+", "Monthly operating costs by hub"),
        ("hub_capacity_jan2026", "5", "Staff, mentors, caseload, burn, throughput headroom"),
        ("funding_by_source", "48", "Quarterly income by funder"),
        ("geographic_waitlist", "24", "Waitlist by district - totals 1,100 (Jan 2026)"),
    ], widths=[2.2, 1.4, 2.8])

    add_heading(doc, "Key fields", 1)
    add_table(doc, ["Field", "Notes"], [
        ("beneficiary_id", "Primary join key; duplicates exist in registry only - use last_updated to resolve"),
        ("last_updated", "ISO date on registry rows; keep the latest when IDs collide"),
        ("placed_90d", "Y/Yes/1/N/No/0/blank; blank = missing follow-up, not 'not placed'"),
        ("completion_status", "Completed / Dropped / Active. Primary 65% rate uses Completed only"),
        ("employment_type", "Formal vs Gig - report both when discussing the 65% threshold"),
        ("province / hub", "Geography of delivery or residence; HCMC vs Ho Chi Minh City"),
        ("date fields", "Enrolment, completion and outcome dates - mixed formats"),
        ("wage_vnd_monthly", "Monthly wage; check notes for the known data-entry error"),
        ("disability_status", "Disability flag where captured"),
        ("volunteer_hours.status", "Active vs Inactive - exclude Inactive for mentor-hour averages unless justified"),
        ("amount_vnd_m", "Funding in millions VND; at least one cell is text with a thousands comma (e.g. 1,250)"),
    ], widths=[2.0, 4.4])

    add_heading(doc, "Treat the export like a real monitoring dump", 1)
    add_body(doc, "Operational exports are rarely clean. Profile every table before you analyse. "
             "Document what you find, how you handled it, and what Workflow A needs to know "
             "before relying on a rate. Do not assume definitions match the language in public "
             "reports. If a figure in a shared document and a figure you calculate disagree, "
             "investigate - do not silently pick the nicer one.")

    add_heading(doc, "Counts, waitlist and definitions", 1)
    add_body(doc, "beneficiary_registry is sized to align with the programme guideline (~800 young "
             "people served in 2025; ~802 rows including two intentional duplicate IDs). D-01 is "
             "the prior-year public report (592 completers in 2024 and the 71% broad placement "
             "story) - do not treat D-01's completer count as the registry row count. "
             "geographic_waitlist is a separate demand snapshot (exactly 1,100) - do not add it "
             "to served totals. Always state the placement definition (broad vs formal; "
             "completion_status = Completed vs including Dropped) when you report a rate.")
    add_body(doc, "Mentor counts will not match across sources without caveats: D-01 reports "
             "organisation-wide volunteer mentors (~95); hub_capacity_jan2026 lists delivery "
             "mentors on each hub roster (sum ~26); volunteer_hours contains every mentor ID "
             "that logged time (including Inactive M-099). Do not force them to reconcile.")

    add_heading(doc, "Capacity when choosing scale scenarios", 1)
    add_body(doc, "Recommendations must fit GBF's mission and Board ambitions, but also its "
             "internal capacity. Use hub_capacity_jan2026 (staff FTE, active mentors, caseload, "
             "monthly burn, max annual throughput) together with hub_costs_2025 and "
             "funding_by_source before endorsing Hold / Moderate / Aggressive growth. "
             "A scenario that clears a funding gap but breaks caseload or mentor headroom is not viable.")
    _save(doc, "WSB_Data_Dictionary", "3_Workflow_B_Operations_and_Analytics", "pdf")


def build_wsb_workbook():
    doc = _doc("Data Analysis Workbook", "Foundation methods and request handling for Workflow B",
               classification="Participant material - Workflow B", ref="MOM-2026-WSB-02")
    add_body(doc, "This workbook covers foundation work (quality + dashboard) and how to run the "
             "A↔B loop: Data Estate Brief → read Demand → Analysis Plan you own → Findings Memo. "
             "It does not prescribe every deep-dive method. When A states a new demand, design "
             "the analysis yourself; ask your Buddy for coaching questions if you are stuck. "
             "You do not need to write code.")

    add_heading(doc, "Suggested order", 1)
    add_number(doc, "Weeks 1-3: dictionary, profile, clean (DB-1 / B1); write Data Estate Brief for A; "
               "then dashboard pivots (DB-2 / B2).")
    add_number(doc, "Week 3: read Demand Brief; write Analysis Plan(s); answer each R-xx with a "
               "Findings Memo before optional deep-dives.")
    add_number(doc, "Funding scenarios (B5) once A has draft options or asks what growth costs.")

    add_heading(doc, "Concepts in plain English", 1)
    add_table(doc, ["Term", "What it means"], [
        ("Pivot table", "A table that groups and summarises your data automatically"),
        ("XLOOKUP", "Finds a value in one table and returns a matching value from another"),
        ("Placement rate", "Placed completers divided by total completers - define both parts"),
        ("Correlation", "Two things move together; it does not prove one causes the other"),
        ("Confounder", "A hidden factor that affects both things you are comparing"),
        ("Selection bias", "Those who attend more may already differ in ways that affect outcomes"),
        ("Data Estate Brief", "Catalog for A of what you hold - never the raw workbook"),
        ("Analysis Plan", "Your method to answer A's demand - you own it; A does not co-edit pivots"),
    ], widths=[1.8, 4.6])

    add_heading(doc, "From Demand to Findings Memo", 1)
    add_number(doc, "Read A's Demand Brief (and later each R-xx). Restate the decision question.")
    add_number(doc, "Write an Analysis Plan (AP-xx): approach, definitions, risks, ETA. File in "
               "A_B_Exchange/. A may clarify scope - not rewrite your joins.")
    add_number(doc, "Decide which tables answer it. Clean only what you need first.")
    add_number(doc, "Compute the result. Note definitions (Completed vs Dropped; Formal vs Broad).")
    add_number(doc, "Write the Findings Memo in plain language. Attach at most two curated charts.")
    add_number(doc, "Update the Request Log status to 'Returned'.")
    add_body(doc, "If Demand is vague ('send us everything on mentoring'), send it back for a "
             "decision question. If the method is unclear, ask Buddy: 'Which tables would you look "
             "at first?' - not 'What is the answer?' Never paste the data dictionary into the "
             "Estate Brief.")

    add_heading(doc, "Foundation: dashboard pivots (DB-2)", 1)
    add_number(doc, "After cleaning, build placement rates by hub, programme and gender.")
    add_number(doc, "State the placement definition on every chart.")
    add_number(doc, "Keep the workbook internal; only curated views leave via Findings Memos.")

    add_heading(doc, "Reporting standards", 1)
    add_bullet(doc, "State the placement definition you used and why.")
    add_bullet(doc, "Never delete an outlier silently - investigate and record the decision.")
    add_bullet(doc, "Give every chart a title that states the finding.")
    add_bullet(doc, "State placement definition; do not mix served totals with waitlist demand.")
    add_bullet(doc, "Never claim causation from a simple association.")
    add_bullet(doc, "Never attach the raw workbook to a Findings Memo.")
    _save(doc, "WSB_Data_Analysis_Workbook", "3_Workflow_B_Operations_and_Analytics", "pdf")


# ===========================================================================
# 4. SHARED TOOLKIT (guides + templates)
# ===========================================================================

def build_storytelling_guide():
    doc = _doc("Presentation & Storytelling Guide", "How to present to the GBF Board",
               classification="Participant material - shared", ref="MOM-2026-ST-01")
    add_heading(doc, "The one rule", 1)
    add_body(doc, "Every slide answers a single question, and its title states the answer. Build "
             "each slide as: headline (the answer), then evidence, then the insight, then the "
             "implication or recommendation.")
    add_heading(doc, "What to avoid", 1)
    add_table(doc, ["Instead of this", "Do this"], [
        ("A slide titled 'Data analysis'", "A slide titled with the insight, e.g. what the evidence forces the Board to face"),
        ("Twelve numbers on one chart", "One chart making one point"),
        ("A decorative stock photo", "White space and one piece of evidence"),
        ("Eight bullet points of recommendations", "One recommendation per slide"),
        ("Ending on 'thank you'", "Ending on the decision you need from the Board"),
    ], widths=[3.2, 3.2])
    add_heading(doc, "An eight-slide structure (matches the board deck template)", 1)
    for i, (h, e) in enumerate([
        ("Your recommendation in one sentence", "The single most important message"),
        ("What is at stake", "Why the Board cannot postpone the decision"),
        ("What the evidence forces you to confront", "One hard finding, cited"),
        ("Where definitions or data change the story", "Measurement honesty without a lecture"),
        ("The options we weighed", "Two or three, with trade-offs"),
        ("Our recommendation", "The choice and its theory of change"),
        ("How it would be delivered", "A ninety-day roadmap"),
        ("The decision we need today", "Two or three specific asks"),
    ], 1):
        add_bullet(doc, f"Slide {i}: {h} - {e}")
    add_body(doc, "Keep detailed methodology, hub comparisons and full financials in an appendix. "
             "Cite Request IDs on any slide that quotes an operational number. Do not open with "
             "a restatement of the engagement letter - open with your recommendation.")
    _save(doc, "Presentation_and_Storytelling_Guide", "4_Shared_Toolkit", "pdf")


def build_roles_guide():
    doc = _doc("Team Roles & Ways of Working", "Shared reference for the whole team",
               classification="Participant material - shared", ref="MOM-2026-TR-01")
    add_body(doc, "Role assignments also appear in 00_Programme_Guideline.pdf (Section 4) so "
             "you can set them in Week 1 without opening this file. This page adds ways of "
             "working and anti-free-riding rules.")
    add_heading(doc, "Suggested roles", 1)
    add_table(doc, ["Role", "Count", "Owns"], [
        ("Engagement Lead", "1", "Timeline, board narrative, final integration, Request Log health"),
        ("Workflow A Lead", "1", "Issue tree, options, trade-off reflection, Analysis Requests"),
        ("Research Lead", "1", "Interview synthesis, credibility matrix, validation"),
        ("Workflow B Lead", "1", "Data dictionary, cleaning, Findings Memos, dashboard"),
        ("Operations Analyst", "1", "Funding scenarios, cost analysis"),
        ("Project & Quality Lead", "1", "Minutes, version control, handoff checklist, file naming"),
    ], widths=[1.9, 0.7, 3.8])
    add_body(doc, "In a team of five, combine the Research Lead with the Workflow A Lead.")
    add_heading(doc, "Preventing free-riding", 1)
    add_bullet(doc, "Every deliverable has exactly one accountable owner in the RACI.")
    add_bullet(doc, "Each member speaks at the board presentation.")
    add_bullet(doc, "A must file requests; B must answer them - neither can do the other's job.")
    add_bullet(doc, "The time tracker is reviewed every Friday; peer evaluation is confidential.")
    add_heading(doc, "Decision-making", 1)
    add_body(doc, "Decide by consensus where you can. Where you cannot, the Engagement Lead "
             "decides and the disagreement is recorded in the decision log. Run the A↔B loop "
             "(Data Estate Brief → Demand Brief → Analysis Plan → R/FM) before gated submissions; "
             "complete Cross_Workflow_Handoff_Checklist.docx with online ticks - Guideline Section 8.")
    add_heading(doc, "Fonts for Vietnamese", 1)
    add_body(doc, "Use Nunito (body) or Lexend (headings) in Word and PowerPoint so Vietnamese "
             "text does not break. Lora is acceptable for long-form narrative if preferred.")
    _save(doc, "Team_Roles_and_Ways_of_Working", "4_Shared_Toolkit", "pdf")


# ---- Templates (delivered as Word) ----

def _template(name, dest, title, builder):
    doc = _doc(title, "Template - complete and export to PDF before submission",
               classification="Participant template")
    builder(doc)
    _save(doc, name, dest, "docx")


def build_templates():
    # Shared
    def charter(doc):
        add_heading(doc, "1. Engagement summary", 2)
        add_meta_table(doc, [
            ("Team name", "________________________________"),
            ("Client", "Generation Bridge Foundation"),
            ("Decision you are advising on (your framing)", "________________________________"),
            ("Board presentation", "Thursday 10 September 2026, 09:00 ICT"),
        ])
        add_heading(doc, "2. Team members and roles", 2)
        add_table(doc, ["Name", "Role", "Workflow"],
                  [("", "Engagement Lead", ""), ("", "Workflow A Lead", "A"),
                   ("", "Workflow B Lead", "B"), ("", "Research Lead", "A"),
                   ("", "Operations Analyst", "B"), ("", "Project & Quality Lead", "")],
                  widths=[2.6, 2.4, 1.4])
        add_heading(doc, "3. Scope", 2)
        add_bullet(doc, "In scope: strategy recommendation, A↔B evidence loop (Estate/Demand/Plan/R/FM), "
                   "validation against transcripts, board presentation and final report.")
        add_bullet(doc, "Out of scope: legal review, full financial audit, implementation.")
        add_bullet(doc, "Data rule: A does not hold raw datasets; B does not share the raw workbook.")
        add_heading(doc, "4. Ways of working", 2)
        add_table(doc, ["Item", "Agreement"], [
            ("Meeting cadence", ""), ("Communication channel", ""),
            ("Decision rule", "Consensus; escalate to Engagement Lead; log disagreements"),
            ("Analysis requests", "Estate → Demand → Plan → min. three R/FM pairs before final options"),
            ("Handoff deadline", "48 hours before any submission"),
        ], widths=[2.2, 4.2])
        add_heading(doc, "5. Online confirmation", 2)
        add_body(doc, "No wet signature required. Each member ticks [x] and types full name + date "
                 "(YYYY-MM-DD) on this Drive Doc or via a comment. Same method as the "
                 "Cross-Workflow Handoff Checklist (Guideline Section 8).")
        for _ in range(3):
            add_online_confirm(doc, "Team member")
    _template("Project_Charter", "4_Shared_Toolkit", "Project Charter", charter)

    def handoff(doc):
        add_heading(doc, "How to use this file", 2)
        add_body(doc, "Location in the package: 4_Shared_Toolkit/Templates/"
                 "Cross_Workflow_Handoff_Checklist.docx. Copy it into your team Drive before "
                 "editing. Full logistics are in 00_Programme_Guideline.pdf Section 8.")
        add_number(doc, "No wet signature and no printing. Confirm online only.")
        add_number(doc, "Fill every [  ] checklist item that applies, then complete the confirmation "
                   "block under your section: tick [x], type full name, type date (YYYY-MM-DD), "
                   "and note the channel (usually 'this Drive Doc').")
        add_number(doc, "Section A = Workflow A reviews B. Section B = Workflow B reviews A. "
                   "Section C = Engagement Lead joint gate.")
        add_number(doc, "When complete, File > Download > PDF Document, name it "
                   "Team_All_Handoff_vN.pdf, store in your team Drive, and submit the link via "
                   "the weekly Google Form when the Master Timeline requires it (W6-02 and before "
                   "gated finals).")
        add_body(doc, "No deliverable may be submitted until both workflows have completed and "
                 "confirmed their sections. The Engagement Lead confirms Section C only if "
                 "A_B_Exchange/ and the Request Log show genuine exchange (not a single dump "
                 "at the end).", bold=True)
        add_heading(doc, "Section A - Workflow A reviews Workflow B", 2)
        for item in [
            "Data Estate Brief is a catalog only - workbook and dictionary were not shared with A.",
            "Analysis Plans exist for answered demands; B owned the method.",
            "Findings Memos answer the question asked, in plain language.",
            "Each DEB / AP / FM-xx PDF is saved in A_B_Exchange/ and logged on the Request Log.",
            "Placement definition (Completed status; Formal vs Broad) is stated.",
            "Caveats are honest; uncomfortable findings are not hidden.",
            "Unanswered or delayed requests are logged with reasons.",
            "No causation is claimed from correlation.",
            "Charts attached to memos are curated, not a pivot dump.",
        ]:
            add_bullet(doc, "[  ]  " + item)
        add_online_confirm(doc, "Workflow A reviewer",
                           hint="Tick [x], type full name + date. No wet signature.")
        add_heading(doc, "Section B - Workflow B reviews Workflow A", 2)
        for item in [
            "Demand Brief ranked decision needs (not a column shopping list).",
            "Every quantitative claim cites a Request ID or Findings Memo ID.",
            "Document cites include section/heading (not code alone).",
            "At least three Analysis Request PDFs are in A_B_Exchange/ with matching Findings Memos.",
            "Strategic options are financially feasible against D-12 and D-03.",
            "Options changed after at least one B challenge (see revision log).",
            "Journey-map pain points link to recommendations and/or Findings cites.",
            "Trade-off reflection names who gains and who loses.",
            "Implementation roadmap has a realistic timeline and cost.",
        ]:
            add_bullet(doc, "[  ]  " + item)
        add_online_confirm(doc, "Workflow B reviewer",
                           hint="Tick [x], type full name + date. No wet signature.")
        add_heading(doc, "Section C - Joint final confirmation", 2)
        add_body(doc, "A_B_Exchange/ shows Data Estate Brief + Demand Brief + Analysis Plan(s) "
                 "and Request Log shows minimum three completed R/FM pairs: [  ] Yes")
        add_body(doc, "Drive link to A_B_Exchange/: ________________________________")
        add_body(doc, "Drive link to this completed Handoff PDF: ________________________________")
        add_online_confirm(doc, "Engagement Lead",
                           hint="Tick [x] only after Sections A and B are done. Type full name + date.")
    _template("Cross_Workflow_Handoff_Checklist", "4_Shared_Toolkit",
              "Cross-Workflow Handoff Checklist", handoff)

    def data_estate_brief(doc):
        add_body(doc, "OFFICIAL CHANNEL - Workflow B → Workflow A. Brief A on what you hold "
                 "without handing over the workbook or data dictionary. Export to PDF as "
                 "Team_WSB_DEB_vN.pdf in A_B_Exchange/. Update if the estate changes materially.")
        add_body(doc, "Allowed: table names, grain (what one row means), time coverage, what each "
                 "table is for, known limits in plain language. Forbidden: field-by-field dumps "
                 "that let A write the analysis for you; row-level extracts; the raw file.",
                 bold=True)
        add_meta_table(doc, [
            ("Brief ID", "DEB-01"),
            ("Date", ""),
            ("Author (B name + role)", ""),
            ("Based on profiling as of", ""),
            ("Saved to A_B_Exchange? (Y/N)", ""),
        ])
        add_heading(doc, "1. Tables we hold (catalog)", 2)
        add_table(doc, ["Table (name only)", "Grain (one row = ...)", "Rough coverage / period",
                        "What it helps answer (plain language)"],
                  [("", "", "", "")] * 9, widths=[1.5, 1.5, 1.4, 2.0])
        add_heading(doc, "2. Joins and keys (plain language)", 2)
        add_body(doc, "How tables connect at a high level (e.g. 'person ID links registry to "
                 "outcomes'). Do not paste a schema dump.")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "3. Known limits A should expect", 2)
        add_body(doc, "Coding messiness, missingness, definition traps - without giving the "
                 "answer to the Board decision.")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "4. What this brief is not", 2)
        add_bullet(doc, "[  ]  We did not attach GBF_Datasets.xlsx or the data dictionary.")
        add_bullet(doc, "[  ]  We did not share pivot-ready extracts or 'just these columns'.")
        add_bullet(doc, "[  ]  We did not pre-answer A's strategy questions in this brief.")
        add_heading(doc, "5. Workflow A acknowledgement", 2)
        add_meta_table(doc, [("Received by (A)", ""), ("Date read", ""),
                             ("Questions for B before Demand Brief", "")])
    _template("Data_Estate_Brief", "4_Shared_Toolkit", "Data Estate Brief", data_estate_brief)

    def demand_brief(doc):
        add_body(doc, "OFFICIAL CHANNEL - Workflow A → Workflow B. Tell B what you need for the "
                 "Board decision after reading the Data Estate Brief. Export to PDF as "
                 "Team_WSA_DMD_vN.pdf in A_B_Exchange/. This seeds R-01+; it is not a column list.")
        add_body(doc, "Demand = decision needs and hypotheses. Do not prescribe field names you "
                 "were not given. Rank what would change your recommendation.", bold=True)
        add_meta_table(doc, [
            ("Brief ID", "DMD-01"),
            ("Date", ""),
            ("Filed by (A name + role)", ""),
            ("Read Data Estate Brief ID", "DEB-01"),
            ("Saved to A_B_Exchange? (Y/N)", ""),
        ])
        add_heading(doc, "1. Decision the Board must make (one sentence)", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "2. Ranked demands", 2)
        add_table(doc, ["Rank", "Decision need / hypothesis", "Why it would change our recommendation",
                        "Useful answer shape (not field names)", "Linked R-xx (when filed)"],
                  [("1", "", "", "", "R-01"),
                   ("2", "", "", "", "R-02"),
                   ("3", "", "", "", "R-03"),
                   ("4", "", "", "", ""),
                   ("5", "", "", "", "")],
                  widths=[0.5, 1.6, 1.6, 1.6, 1.0])
        add_heading(doc, "3. Qualitative sources that created these demands", 2)
        add_body(doc, "(Transcript / document cites with section or speaker)")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "4. What we are not asking (out of scope for now)", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "5. Workflow B acknowledgement", 2)
        add_meta_table(doc, [("Received by (B)", ""), ("Analysis Plan ID(s)", ""),
                             ("Target plan date", "")])
    _template("Demand_Brief", "4_Shared_Toolkit", "Demand Brief", demand_brief)

    def analysis_plan(doc):
        add_body(doc, "OFFICIAL CHANNEL - Workflow B. You own the method. After A's Demand Brief "
                 "(and for each formal R-xx as needed), write how you will answer. Export to PDF "
                 "as Team_WSB_APxx_Plan_vN.pdf in A_B_Exchange/. Share the plan with A; do not "
                 "share interim raw extracts.")
        add_body(doc, "A may ask clarifying questions about scope and timing. A does not co-edit "
                 "your pivots or rewrite your join path.", bold=True)
        add_meta_table(doc, [
            ("Plan ID (AP-xx)", ""),
            ("Responds to Demand / Request", "DMD-01 item __ / R-__"),
            ("Date", ""),
            ("Author (B name + role)", ""),
            ("Saved to A_B_Exchange? (Y/N)", ""),
        ])
        add_heading(doc, "1. Decision question restated (B's words)", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "2. Approach (B owns this)", 2)
        add_body(doc, "Which tables you expect to use, how you will define key measures, what you "
                 "will clean or exclude, and what comparison you will produce. Keep it readable "
                 "for A - not a SQL dump.")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "3. Definitions you will use", 2)
        add_body(doc, "(Especially placement / completion / population)")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "4. Risks and what this plan cannot prove", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "5. Timeline and Findings Memo target", 2)
        add_meta_table(doc, [("Work start", ""), ("FM-xx target date", ""),
                             ("Owner", ""), ("Buddy informed? (Y/N)", "")])
        add_heading(doc, "6. Workflow A acknowledgement", 2)
        add_meta_table(doc, [("Read by (A)", ""), ("Date", ""),
                             ("Scope OK / questions", "")])
    _template("Analysis_Plan", "4_Shared_Toolkit", "Analysis Plan", analysis_plan)

    def analysis_request(doc):
        add_body(doc, "OFFICIAL CHANNEL - Workflow A. Complete one form per ask. Export to PDF and "
                 "save in your team Drive folder A_B_Exchange/ using naming "
                 "Team_WSA_Rxx_Request_vN.pdf. Also log the request on the Request Log tab. "
                 "Chat or verbal asks do not count for assessment.")
        add_body(doc, "File R-xx after the Data Estate Brief, Demand Brief and (where ready) "
                 "Analysis Plan exist. Do not prescribe field names you have not been given. "
                 "Working out what to ask is part of the exercise.")
        add_meta_table(doc, [
            ("Request ID (R-xx)", ""),
            ("Linked Demand Brief item", "DMD-01 rank __"),
            ("Linked Analysis Plan (AP-xx)", ""),
            ("Date filed", ""),
            ("Filed by (name + role)", ""),
            ("Urgency (routine / needed for options / blocking)", ""),
            ("Saved to A_B_Exchange? (Y/N)", ""),
        ])
        add_heading(doc, "1. Decision this informs", 2)
        add_body(doc, "What choice would this answer change?")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "2. Hypothesis or claim to test", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "3. What a useful answer would look like", 2)
        add_body(doc, "(For example: a comparison across hubs; a rate under a stated definition; "
                 "a cost range. Be specific about the decision, not the spreadsheet.)")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "4. Qualitative sources that prompted this ask", 2)
        add_body(doc, "(Transcript or document refs with section/speaker)")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "5. What would change your mind", 2)
        add_body(doc, "If the operational record showed X instead of Y, we would ________________")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "6. Workflow B acknowledgement", 2)
        add_meta_table(doc, [("Received by", ""), ("Target return date", ""),
                             ("Findings Memo ID (FM-xx)", "")])
    _template("Analysis_Request", "4_Shared_Toolkit", "Analysis Request", analysis_request)

    def findings_memo(doc):
        add_body(doc, "OFFICIAL CHANNEL - Workflow B. Complete one memo per Analysis Request. "
                 "Export to PDF and save in A_B_Exchange/ as Team_WSB_FMxx_Findings_vN.pdf. "
                 "Update the Request Log to Returned. Do not attach the raw workbook. Write for "
                 "a non-analyst.")
        add_meta_table(doc, [
            ("Findings Memo ID (FM-xx)", ""),
            ("Responds to Request ID", ""),
            ("Follows Analysis Plan", "AP-__"),
            ("Date returned", ""),
            ("Author (name + role)", ""),
            ("Saved to A_B_Exchange? (Y/N)", ""),
        ])
        add_heading(doc, "1. Question restated", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "2. Method vs plan (brief)", 2)
        add_body(doc, "What you did relative to AP-xx; cleaning choices that affect this answer; "
                 "definition note:")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "3. Finding in plain language", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "4. Definitions used", 2)
        add_body(doc, "(Especially placement / completion)")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "5. Caveats and what this does not prove", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "6. Implication for A's decision (optional one sentence)", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "7. Attachments", 2)
        add_body(doc, "List curated charts or summary tables only (max two recommended). "
                 "File names of any charts also saved in A_B_Exchange/:")
        add_body(doc, "________________________________________________________________")
        add_body(doc, "Linked trail: Demand item __ / AP-__ / R-__ (fill all that apply).",
                 italic=True, size=10)
    _template("Findings_Memo", "4_Shared_Toolkit", "Findings Memo", findings_memo)

    def minutes(doc):
        add_meta_table(doc, [("Meeting", ""), ("Date", ""), ("Time", ""),
                             ("Attendees", ""), ("Absent", ""), ("Note-taker", "")])
        add_heading(doc, "Decisions", 2)
        add_table(doc, ["Decision", "Rationale", "Owner"], [("", "", ""), ("", "", "")],
                  widths=[2.6, 2.6, 1.2])
        add_heading(doc, "Actions", 2)
        add_table(doc, ["Action", "Owner", "Due", "Status"],
                  [("", "", "", ""), ("", "", "", "")], widths=[2.8, 1.4, 1.2, 1.0])
        add_heading(doc, "Open Analysis Requests", 2)
        add_table(doc, ["Request ID", "Status", "Owner"], [("", "", ""), ("", "", "")],
                  widths=[1.4, 2.6, 2.4])
        add_body(doc, "Minutes are circulated within 24 hours; corrections raised within 48 hours.")
    _template("Meeting_Minutes", "4_Shared_Toolkit", "Meeting Minutes", minutes)

    def execsum(doc):
        add_heading(doc, "Key message", 2)
        add_body(doc, "[One sentence: the single most important thing the Board should know.]")
        add_heading(doc, "Context", 2)
        add_body(doc, "[Three or four sentences on why this decision matters now.]")
        add_heading(doc, "Recommendation", 2)
        add_number(doc, "Primary recommendation:")
        add_number(doc, "Supporting actions:")
        add_number(doc, "What GBF should stop doing:")
        add_heading(doc, "Evidence", 2)
        add_table(doc, ["Finding", "Source (D-/ST-/R-/FM-)", "Implication"],
                  [("", "", ""), ("", "", "")], widths=[2.4, 1.8, 2.2])
        add_heading(doc, "Ninety-day roadmap", 2)
        add_table(doc, ["Phase", "Actions", "Cost estimate"],
                  [("0-30 days", "", ""), ("31-60 days", "", ""), ("61-90 days", "", "")],
                  widths=[1.4, 3.4, 1.6])
        add_heading(doc, "What we still do not know", 2)
        add_body(doc, "[Honest limitations - this builds credibility.]")
    _template("Executive_Summary", "4_Shared_Toolkit", "Executive Summary", execsum)

    def finalrep(doc):
        add_body(doc, "Recommended length 15-25 pages excluding the appendix. Use this structure.")
        for sec in [
            "Executive summary", "Engagement approach (including A↔B evidence loop)",
            "Problem definition", "Situation analysis", "Stakeholder insights",
            "Findings from Analysis Requests", "Strategic options", "Recommendation",
            "Implementation roadmap", "Trade-offs accepted", "Risks and mitigations",
            "Limitations and further work", "Appendices (Request Log extract, key Findings Memos)",
        ]:
            add_number(doc, sec)
        add_heading(doc, "Standards", 2)
        add_bullet(doc, "Every claim carries a source reference (document, transcript, or R/FM ID).")
        add_bullet(doc, "The placement definition is stated wherever the rate is used.")
        add_bullet(doc, "Charts are sourced and dated.")
    _template("Final_Report", "4_Shared_Toolkit", "Final Report", finalrep)

    def peer(doc):
        add_body(doc, "Confidential. Submit individually by Monday 14 September 2026, 17:00 ICT. "
                 "Facilitators use aggregate results; individual ratings remain confidential.")
        add_body(doc, "Rate each teammate from 1 (poor) to 5 (excellent).")
        add_table(doc, ["Team member", "Contribution", "Quality", "Communication",
                        "Reliability", "Collaboration"],
                  [("", "", "", "", "", ""), ("", "", "", "", "", ""),
                   ("", "", "", "", "", ""), ("", "", "", "", "", "")],
                  widths=[1.7, 1.0, 0.9, 1.2, 1.0, 1.1])
        add_heading(doc, "Open questions", 2)
        add_body(doc, "1. Who contributed most to the team's success, and why?")
        add_body(doc, "2. Did the Analysis Request / Findings Memo exchange work in practice?")
        add_body(doc, "3. Anything the facilitators should know (optional)?")
    _template("Peer_Evaluation", "4_Shared_Toolkit", "Peer Evaluation (Confidential)", peer)

    def reflect(doc):
        add_body(doc, "Individual and private. Submit by Tuesday 15 September 2026, 17:00 ICT.")
        for q in ["What did you learn about asking for evidence rather than assuming it?",
                  "What was hardest about working across two workflows with incomplete information?",
                  "Describe one trade-off your team faced and what you decided.",
                  "What did you contribute, and what would you do differently?",
                  "One skill you want to develop next:"]:
            add_heading(doc, q, 2)
            add_body(doc, "________________________________________________________________")
            add_body(doc, "________________________________________________________________")
    _template("Reflection_Journal", "4_Shared_Toolkit", "Reflection Journal (Individual)", reflect)

    # Workflow A templates
    def issue_tree(doc):
        add_heading(doc, "Problem / decision statement", 2)
        add_body(doc, "In one precise sentence, state the decision the Board must make. Use your "
                 "own words. Copying or lightly editing the engagement letter is not enough.")
        add_body(doc, "________________________________________________________________")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "Issue tree (MECE)", 2)
        add_body(doc, "Invent the top-level branches from the evidence in the data room. "
                 "Sub-questions at each level must be mutually exclusive and collectively "
                 "exhaustive. Do not force the tree into a pre-set list of topics.")
        add_table(doc, ["Branch (your label)", "Sub-questions", "Testable hypothesis", "Sources"], [
            ("", "", "", ""), ("", "", "", ""), ("", "", "", ""),
            ("", "", "", ""), ("", "", "", "")], widths=[1.4, 1.7, 1.7, 1.6])
        add_body(doc, "Sources column: cite transcript CODE (Speaker) and/or document Ref + "
                 "section (e.g. D-02 §1; BN-02 Tuan). Do not leave Sources blank for any "
                 "hypothesis you intend to test.", italic=True, size=10)
        add_heading(doc, "Hypotheses to test (each should imply an Analysis Request)", 2)
        for i in range(1, 5):
            add_body(doc, f"H{i}: ________________________________________________")
            add_body(doc, f"    Sources: ________________________________________")
            add_body(doc, f"    Possible ask to Workflow B: ________________________")
    _template("Problem_Statement_and_Issue_Tree", "2_Workflow_A_Impact_Strategy",
              "Problem Statement and Issue Tree", issue_tree)

    def stakeholder(doc):
        add_body(doc, "Plot stakeholders on a power/interest grid. Star the five you must "
                 "understand most deeply. Cite at least one source per starred stakeholder.")
        add_table(doc, ["Stakeholder", "Power (1-5)", "Interest (1-5)", "Position on strategy", "Source"],
                  [("", "", "", "", "")] * 8, widths=[1.8, 1.0, 1.0, 2.0, 1.6])
    _template("Stakeholder_Map", "2_Workflow_A_Impact_Strategy", "Stakeholder Map", stakeholder)

    def credibility(doc):
        add_body(doc, "You rate each shared source - the data room does not do it for you. "
                 "For each row, write why you trust or doubt it (author, audience, incentive, "
                 "definitions, contradictions with other sources). Operational datasets are held "
                 "by Workflow B; rate Findings Memos once received.")
        add_table(doc, ["Source", "Reliability (H/M/L)", "Why", "Key fact you trust or doubt"],
                  [("D-01 Annual Report", "", "", ""), ("D-02 Internal memo", "", "", ""),
                   ("D-03 VPBank agreement", "", "", ""), ("D-08 Board emails", "", "", ""),
                   ("D-10 YouthWorks brief", "", "", ""), ("D-11 Staff survey", "", "", ""),
                   ("D-12 Financials", "", "", ""), ("Interview transcripts", "", "", ""),
                   ("Findings Memos (from B)", "", "", "")],
                  widths=[1.6, 1.2, 2.0, 2.6])
        add_body(doc, "Prompt questions (optional): Who benefits if this source is believed? "
                 "What definition sits behind any rate? What would change your rating?")
    _template("Information_Credibility_Matrix", "4_Shared_Toolkit",
              "Information Credibility Matrix", credibility)

    def synth(doc):
        add_body(doc, "Summarise by theme, not by person. Every insight cites a transcript.")
        add_table(doc, ["Theme", "Insight", "Sources", "Contradicts?"],
                  [("", "", "", "")] * 5, widths=[1.5, 2.6, 1.2, 1.1])
        add_heading(doc, "Contradictions to resolve", 2)
        add_body(doc, "For example: is YouthWorks a serious threat? Compare ST-04, ST-05 and D-10.")
        add_heading(doc, "Questions this raises for Workflow B", 2)
        add_body(doc, "________________________________________________________________")
    _template("Interview_Synthesis", "2_Workflow_A_Impact_Strategy",
              "Interview Synthesis", synth)

    def empathy(doc):
        add_heading(doc, "Chosen archetype", 2)
        add_body(doc, "Ground this in a specific voice from the transcripts (cite BN-xx).")
        add_heading(doc, "Empathy map (cite quotes)", 2)
        add_table(doc, ["Says", "Thinks"], [("", "")], widths=[3.2, 3.2])
        add_table(doc, ["Does", "Feels"], [("", "")], widths=[3.2, 3.2])
        add_heading(doc, "Journey map", 2)
        add_table(doc, ["Stage", "What happens", "Emotion", "Pain point", "Evidence (quote or FM-)"], [
            ("Awareness", "", "", "", ""), ("Waitlist", "", "", "", ""),
            ("Training", "", "", "", ""), ("Internship", "", "", "", ""),
            ("Job search", "", "", "", ""), ("90-day outcome", "", "", "", "")],
            widths=[1.1, 1.4, 0.9, 1.2, 1.8])
    _template("Empathy_and_Journey_Map", "2_Workflow_A_Impact_Strategy",
              "Empathy & Journey Map", empathy)

    def options(doc):
        add_body(doc, "Present two or three genuinely different options. Cite Request / Findings "
                 "IDs for every quantitative claim.")
        for letter in ["A", "B", "C"]:
            add_heading(doc, f"Option {letter}", 2)
            add_meta_table(doc, [("Description", ""), ("Rationale", ""),
                                 ("Cost implication (cite FM-/D-12)", ""),
                                 ("Main risks", ""), ("Who wins / who loses", ""),
                                 ("Evidence cites (R-/FM-/ST-/D-)", "")])
        add_heading(doc, "Recommendation", 2)
        add_body(doc, "We recommend Option ____ because ____________________________.")
        add_body(doc, "We would change our mind if ____________________________.")
    _template("Strategic_Options", "2_Workflow_A_Impact_Strategy",
              "Strategic Options", options)

    def tradeoff(doc):
        add_body(doc, "Complete after you have a preferred recommendation. Answer in your own "
                 "words. There is no checklist of named dilemmas to tick.")
        add_heading(doc, "1. Who gains under your recommendation?", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "2. Who loses, waits longer, or is deprioritised?", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "3. What would you need to believe for the opposite choice to be right?", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "4. What cost does your preferred path accept, and why is that acceptable?", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "5. Which transcript voice would be hardest to explain this to - and how "
                    "would you explain it?", 2)
        add_body(doc, "________________________________________________________________")
    _template("Tradeoff_Reflection", "2_Workflow_A_Impact_Strategy",
              "Trade-off Reflection", tradeoff)

    def validation(doc):
        add_body(doc, "Use interview transcripts and Findings Memos. Map each draft recommendation "
                 "to evidence. Minimum six citations: at least two beneficiary (BN-), two staff "
                 "(ST-), one donor (DN-) and one employer (EP-). Include at least two R-/FM- cites.")
        add_table(doc, ["Recommendation", "Source", "What it says", "Supports / challenges", "Action"],
                  [("", "", "", "", "")] * 6, widths=[1.5, 1.0, 1.6, 1.3, 1.0])
    _template("Validation_Protocol", "2_Workflow_A_Impact_Strategy",
              "Validation Protocol", validation)

    def revlog(doc):
        add_body(doc, "Record how validation and Findings Memos changed your thinking. At least "
                 "one substantive revision after a Workflow B challenge is expected.")
        add_table(doc, ["Original element", "What challenge revealed", "Revised?", "New position"],
                  [("", "", "", "")] * 5, widths=[1.8, 2.2, 0.9, 1.5])
    _template("Recommendation_Revision_Log", "2_Workflow_A_Impact_Strategy",
              "Recommendation Revision Log", revlog)

    # Workflow B templates
    def dqr(doc):
        add_body(doc, "Internal to Workflow B. Two pages maximum. Share with A only material "
                 "issues that affect a decision, via Findings Memo.")
        add_heading(doc, "Issues identified", 2)
        add_table(doc, ["#", "Table", "Issue", "Severity", "Action taken"],
                  [("", "", "", "", "")] * 8, widths=[0.4, 1.6, 2.0, 1.0, 1.4])
        add_heading(doc, "Cleaning decisions and rationale", 2)
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "Placement definition used", 2)
        add_body(doc, "State which definition you adopted and why.")
        add_heading(doc, "What Workflow A needs to know (for FM translation)", 2)
        add_body(doc, "________________________________________________________________")
    _template("Data_Quality_Report", "3_Workflow_B_Operations_and_Analytics",
              "Data Quality Report", dqr)

    def analysis(doc):
        add_heading(doc, "Headline findings", 2)
        add_table(doc, ["Finding", "Linked Request / FM", "Caveat"], [("", "", "")] * 4,
                  widths=[2.4, 1.8, 2.2])
        add_heading(doc, "Deep-dives completed (only if A requested them)", 2)
        add_body(doc, "Summarise any on-request analyses (e.g. placement drivers, attendance, "
                 "mentoring, completion). For each: what A asked, what you found, definitions used, "
                 "and whether you are claiming association or causation.")
        add_body(doc, "________________________________________________________________")
        add_heading(doc, "Open questions for further requests", 2)
        add_body(doc, "________________________________________________________________")
    _template("Analysis_Summary", "3_Workflow_B_Operations_and_Analytics",
              "Analysis Summary", analysis)

    def evcheck(doc):
        add_body(doc, "Workflow B checks each Workflow A option against Findings Memos and D-12/D-03.")
        add_table(doc, ["Option / claim", "Supported (cite FM-)?", "Financially feasible?", "Concern"],
                  [("", "", "", "")] * 5, widths=[2.0, 1.6, 1.4, 1.4])
    _template("Evidence_Check_Memo", "3_Workflow_B_Operations_and_Analytics",
              "Evidence-Check Memo", evcheck)


def build_all():
    print("Building realistic documents...")
    build_programme_handbook()
    build_data_room_index()
    build_deadlines()
    build_d01_annual_report()
    build_d02_programme_review()
    build_d03_grant_agreement()
    build_d08_email_thread()
    build_d10_competitor()
    build_d11_staff_survey()
    build_d12_financials()
    build_interviews()
    build_wsa_start()
    build_wsb_start()
    build_wsa_brief()
    build_wsa_playbook()
    build_wsb_brief()
    build_wsb_data_dictionary()
    build_wsb_workbook()
    build_storytelling_guide()
    build_roles_guide()
    build_templates()
    return MANIFEST


if __name__ == "__main__":
    build_all()
