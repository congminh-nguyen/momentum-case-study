"""Shared Momentum Programme branding and document helpers.

Templates use Nunito (body) and Lexend (headings) so Vietnamese text
renders without broken fonts when Momentees type in Vietnamese.
"""

from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Momentum Programme brand palette
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
TEAL = RGBColor(0x00, 0x7A, 0x87)
ORANGE = RGBColor(0xE8, 0x5D, 0x04)
GREY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x22, 0x22, 0x22)

# Vietnamese-safe stack (install Nunito / Lexend / Lora if missing locally)
FONT = "Nunito"
FONT_HEADING = "Lexend"
FONT_FALLBACK = "Calibri"
PROGRAMME = "Momentum Programme"
CASE_TITLE = "Scaling with Integrity"
CLIENT = "Generation Bridge Foundation"
VERSION = "3.0 | July 2026"


def _set_run_font(run, name=FONT, east_asia=None):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    ea = east_asia or name
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), ea)
    rFonts.set(qn("w:cs"), name)


def set_doc_defaults(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    # Ensure eastAsia for Vietnamese in Normal style
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_cover(doc, title, subtitle="", classification="Participant material",
              doc_ref="", issued=""):
    band = doc.add_paragraph()
    r = band.add_run(PROGRAMME + "   |   " + CASE_TITLE)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = TEAL
    _set_run_font(r, FONT_HEADING)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = NAVY
    _set_run_font(r, FONT_HEADING)

    if subtitle:
        p = doc.add_paragraph()
        r = p.add_run(subtitle)
        r.italic = True
        r.font.size = Pt(12)
        r.font.color.rgb = GREY
        _set_run_font(r, FONT)

    meta_rows = [("Client", CLIENT)]
    if doc_ref:
        meta_rows.append(("Document reference", doc_ref))
    if issued:
        meta_rows.append(("Issued", issued))
    meta_rows.append(("Classification", classification))
    meta_rows.append(("Version", VERSION))
    add_meta_table(doc, meta_rows)
    _rule(doc)


def _rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "007A87")
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        _set_run_font(run, FONT_HEADING)
    return h


def add_body(doc, text, bold=False, italic=False, size=11, color=DARK, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    _set_run_font(r, FONT)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    r = p.add_run(text)
    _set_run_font(r, FONT)
    return p


def add_number(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    r = p.add_run(text)
    _set_run_font(r, FONT)
    return p


def add_quote(doc, text, attribution=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.3)
    r = p.add_run(text if text.startswith("\u201c") else f"\u201c{text}\u201d")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x33, 0x3B, 0x4A)
    _set_run_font(r, FONT)
    if attribution:
        a = doc.add_paragraph()
        a.paragraph_format.left_indent = Inches(0.4)
        ar = a.add_run("\u2014 " + attribution)
        ar.font.size = Pt(10)
        ar.font.color.rgb = GREY
        _set_run_font(ar, FONT)
    return p


def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.allow_autofit = True
    for ri, (label, value) in enumerate(rows):
        c0 = table.rows[ri].cells[0]
        c1 = table.rows[ri].cells[1]
        c0.width = Inches(1.8)
        c1.width = Inches(4.6)
        r0 = c0.paragraphs[0].add_run(str(label))
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = NAVY
        _set_run_font(r0, FONT)
        r1 = c1.paragraphs[0].add_run(str(value))
        r1.font.size = Pt(10)
        r1.font.color.rgb = DARK
        _set_run_font(r1, FONT)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade_cell(hdr[i], "1B2A4A")
        para = hdr[i].paragraphs[0]
        run = para.add_run(str(h))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_run_font(run, FONT_HEADING)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            para = cells[ci].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(10)
            _set_run_font(run, FONT)
            if ri % 2 == 1:
                _shade_cell(cells[ci], "F2F4F7")
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_signature_block(doc, name, role, date_label="Date"):
    """Legacy wet-ink block - prefer add_online_confirm for remote teams."""
    doc.add_paragraph()
    add_body(doc, "Signed:  ______________________________________", size=10)
    add_body(doc, f"{name}", bold=True, size=10, space_after=0)
    add_body(doc, role, italic=True, size=10, color=GREY, space_after=0)
    add_body(doc, f"{date_label}:  ____________________", size=10)


def add_online_confirm(doc, role_label, hint="Tick [x] and type full name + date (online confirmation)"):
    """Online-friendly confirmation instead of wet signature."""
    add_body(doc, f"[  ]  Confirmed by {role_label}", bold=True, size=10, space_after=2)
    add_body(doc, hint, italic=True, size=9, color=GREY, space_after=2)
    add_body(doc, "Name: ____________________    Date: __________    Channel (Drive comment / sheet tick): __________",
             size=10)


def add_page_break(doc):
    doc.add_page_break()
